import os
import pandas as pd
from docx import Document
import re
from datetime import datetime

class ExcelToWordConverter:
    def __init__(self, template_path, output_dir="output"):
        """Initialize the converter with the template file path."""
        self.template_path = template_path
        self.output_dir = output_dir
        
        # Create output directory if it doesn't exist
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
    
    def _parse_filename(self, filename):
        """Parse Excel filename to extract term, std, subject info."""
        # Example: "TERM1_XI_ECO.xlsx"
        parts = os.path.basename(filename).split('_')
        if len(parts) >= 3:
            term = parts[0]
            std = parts[1]  # XI or XII
            subject = parts[2].split('.')[0]  # Remove file extension
            return term, std, subject
        return None, None, None
    
    def _std_label_map(self, std):
        """Map the standard to its full form."""
        std_map = {
            "XI": "FYJC",
            "XII": "SYJC"
        }
        return std_map.get(std, std)
    
    def _get_month_number(self, month_name):
        """Convert month name to its number."""
        month_dict = {
            "JANUARY": "01", "FEBRUARY": "02", "MARCH": "03", "APRIL": "04",
            "MAY": "05", "JUNE": "06", "JULY": "07", "AUGUST": "08",
            "SEPTEMBER": "09", "OCTOBER": "10", "NOVEMBER": "11", "DECEMBER": "12"
        }
        return month_dict.get(month_name.upper(), "")
    
    def _determine_term_month(self, term, month):
        """Determine which month number it is in the term."""
        term1_months = ["JUNE", "JULY", "AUGUST", "SEPTEMBER", "OCTOBER"]
        term2_months = ["NOVEMBER", "DECEMBER", "JANUARY", "FEBRUARY", "MARCH"]
        
        if term == "TERM1":
            try:
                return f"{term1_months.index(month.upper()) + 1:02d}"
            except ValueError:
                return "01"
        elif term == "TERM2":
            try:
                return f"{term2_months.index(month.upper()) + 1:02d}"
            except ValueError:
                return "01"
        return "01"
    
    def convert_excel_to_word(self, excel_path, month):
        """
        Convert data from Excel to Word document.
        
        Args:
            excel_path: Path to the Excel file
            month: Month to filter data by
        """
        # Parse the filename to get term, std, subject
        term, std, subject = self._parse_filename(excel_path)
        if not all([term, std, subject]):
            print(f"Could not parse filename: {excel_path}")
            return
        
        # Load Excel data
        try:
            df = pd.read_excel(excel_path)
        except Exception as e:
            print(f"Error reading Excel file: {e}")
            return
        
        # Filter data for the specified month
        if 'Month' in df.columns:
            df_month = df[df['Month'].str.upper() == month.upper()]
        else:
            print("No 'Month' column found in Excel file")
            return
        
        if df_month.empty:
            print(f"No data found for month: {month}")
            return
        
        # Load the Word template
        try:
            doc = Document(self.template_path)
        except Exception as e:
            print(f"Error loading Word template: {e}")
            return
        
        # Get current academic year
        current_year = datetime.now().year
        if datetime.now().month < 4:  # If before April, it's the previous academic year
            academic_year = f"{current_year-1}-{current_year}"
        else:
            academic_year = f"{current_year}-{current_year+1}"
        
        # Replace placeholders in the document
        for paragraph in doc.paragraphs:
            if 'MONTH' in paragraph.text:
                paragraph.text = paragraph.text.replace('MONTH', month.upper())
            
            # Replace year and reference number placeholder
            month_num = self._get_month_number(month)
            term_month_num = self._determine_term_month(term, month)
            year_pattern = re.compile(r'(\d{4}-\d{4}-\d{2}/\s*[A-Z]\s*–\s*[A-Z]{2}/[A-Z]{3}/\d{2})')
            
            match = year_pattern.search(paragraph.text)
            if match:
                replacement = f"{academic_year}-{month_num}/{subject[0]}–{subject}/{std}/{term_month_num}"
                paragraph.text = year_pattern.sub(replacement, paragraph.text)
        
        # Fill the table with data
        for table in doc.tables:
            # Assuming the first row is headers and first column is for rows
            if len(table.rows) > 1:
                for i, row in enumerate(table.rows[1:], 1):
                    if i <= len(df_month):
                        data_row = df_month.iloc[i-1]
                        
                        # Fill in the cells based on the headers and data
                        for j, cell in enumerate(row.cells):
                            if j == 0:  # First column - index or serial number
                                cell.text = str(i)
                            elif j == 1:  # Initials column
                                cell.text = str(data_row.get('Initials', ''))
                            elif j == 2:  # Standard column
                                std_value = str(data_row.get('STD', ''))
                                cell.text = self._std_label_map(std_value)
                            elif j == 3:  # Allotted column
                                cell.text = str(data_row.get('Allotted', ''))
                            elif j == 4:  # E-Act column mapped to ENGAGED
                                cell.text = str(data_row.get('E-Act', ''))
                            elif j == 5:  # E-Add column mapped to GAP
                                cell.text = str(data_row.get('E-Add', ''))
        
        # Save the document
        output_filename = f"{std}_{subject}_{month}_{term}.docx"
        output_path = os.path.join(self.output_dir, output_filename)
        doc.save(output_path)
        print(f"Created Word document: {output_path}")
        return output_path

def main():
    # Example usage
    template_path = "executive_summary_template.docx"
    converter = ExcelToWordConverter(template_path)
    
    # Convert a specific Excel file for a specific month
    excel_file = "iso_excel.xlsx"  # Change to your actual file
    month = "JUNE"  # Change to the month you want to process
    
    if os.path.exists(excel_file):
        output_path = converter.convert_excel_to_word(excel_file, month)
        if output_path and os.path.exists(output_path):
            print(f"Successfully created {output_path}")
    else:
        print(f"Excel file not found: {excel_file}")

if __name__ == "__main__":
    main()
