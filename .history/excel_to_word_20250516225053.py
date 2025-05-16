import os
import csv
import re
import logging
from docx import Document
import shutil
from copy import deepcopy

# Add win32com import for Word automation
try:
    import win32com.client
    WIN32COM_AVAILABLE = True
except ImportError:
    WIN32COM_AVAILABLE = False

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def parse_filename(filename):
    """Extract information from the filename."""
    # Expected format: iso_excel_YYYY-YYYY_termN_STD.csv
    pattern = r'iso_excel_(\d{4}-\d{4})_term(\d+)_(\w+)\.csv'
    match = re.match(pattern, os.path.basename(filename))
    
    if match:
        year_range, term, std = match.groups()
        # Convert STD to actual standard (FYJC=XI and SYJC=XII)
        std_mapping = {"FYJC": "XI", "SYJC": "XII"}
        standard = std_mapping.get(std, std)
        
        return {
            'year_range': year_range,
            'term': term,
            'standard': standard,
            'original_std': std
        }
    return None

def get_month_number(month_name):
    """Convert month name to number."""
    month_mapping = {
        'JUNE': '06', 'JULY': '07', 'AUG': '08', 'SEPTEMBER': '09', 
        'SEP': '09', 'OCTOBER': '10', 'OCT': '10', 'NOVEMBER': '11',
        'NOV': '11', 'DECEMBER': '12', 'DEC': '12', 'JANUARY': '01',
        'JAN': '01', 'FEBRUARY': '02', 'FEB': '02', 'MARCH': '03',
        'MAR': '03', 'APRIL': '04', 'APR': '04', 'MAY': '05'
    }
    return month_mapping.get(month_name.upper(), '00')

def get_term_month_index(month_name, term):
    """Get the month index within the term."""
    if term == '1':
        months = ['JUNE', 'JULY', 'AUG', 'SEP', 'OCT']
    else:  # term 2
        months = ['NOV', 'DEC', 'JAN', 'FEB']
    
    for i, m in enumerate(months, 1):
        if month_name.upper().startswith(m):
            return f"{i:02d}"
    return "01"  # Default

def replace_placeholders_in_paragraph(paragraph, replacements):
    """Replace placeholders in paragraph text without affecting other formatting."""
    if not any(placeholder in paragraph.text for placeholder in replacements.keys()):
        return False
    
    # Create a new paragraph text with replacements
    text = paragraph.text
    for placeholder, value in replacements.items():
        if placeholder in text:
            text = text.replace(placeholder, value)
    
    # Replace text in the paragraph
    paragraph.text = text
    return True

def find_standard_columns(table):
    """Find the column indices for XI and XII (ALLOTTED, ENGAGED, GAP) by header text."""
    xi_cols = {'ALLOTTED': None, 'ENGAGED': None, 'GAP': None}
    xii_cols = {'ALLOTTED': None, 'ENGAGED': None, 'GAP': None}
    xi_start, xii_start = None, None
    for row in table.rows[:3]:
        for idx, cell in enumerate(row.cells):
            txt = cell.text.strip().upper()
            if txt == 'XI':
                xi_start = idx
            elif txt == 'XII':
                xii_start = idx
    for row in table.rows[:4]:
        for idx, cell in enumerate(row.cells):
            txt = cell.text.strip().upper()
            if xi_start is not None and xi_start <= idx < xi_start + 3:
                if txt == 'ALLOTTED':
                    xi_cols['ALLOTTED'] = idx
                elif txt == 'ENGAGED':
                    xi_cols['ENGAGED'] = idx
                elif txt == 'GAP':
                    xi_cols['GAP'] = idx
            if xii_start is not None and xii_start <= idx < xii_start + 3:
                if txt == 'ALLOTTED':
                    xii_cols['ALLOTTED'] = idx
                elif txt == 'ENGAGED':
                    xii_cols['ENGAGED'] = idx
                elif txt == 'GAP':
                    xii_cols['GAP'] = idx
    logger.info(f"XI columns: {xi_cols}, XII columns: {xii_cols}")
    return xi_cols, xii_cols

def get_placeholder_col_map_and_row(table):
    """Return (placeholder_map, row_index) for the row with all placeholders."""
    for row_idx, row in enumerate(table.rows):
        placeholders = {}
        all_placeholders = True
        for idx, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt.startswith('{{col_') and txt.endswith('}}'):
                placeholders[txt] = idx
            else:
                all_placeholders = False
        if all_placeholders and placeholders:
            logger.info(f"Found placeholder row at index {row_idx}: {placeholders}")
            return placeholders, row_idx
    logger.error("No placeholder row found in table!")
    return None, None

def parse_month_field_columns(header_row, field_row):
    """Return a dict: {month: {field: col_idx}} for each month and field in the CSV."""
    month_field_map = {}
    current_month = None
    month_start_idx = None
    
    # Skip the first two columns (SR.NO. and INITIALS)
    for idx, (month_cell, field_cell) in enumerate(zip(header_row[2:], field_row[2:]), start=2):
        # If the month cell is not empty, update current_month
        if month_cell.strip():
            current_month = month_cell.strip().upper()
            if current_month not in month_field_map:
                month_field_map[current_month] = {}
                month_start_idx = idx
        
        # Only map if we have a current_month and a valid field
        if current_month and month_start_idx is not None:
            field = field_cell.strip().upper()
            if field == 'ALOTTED':
                month_field_map[current_month]['ALLOTTED'] = idx
            elif field == 'ENGAGED':
                month_field_map[current_month]['ENGAGED'] = idx
            elif field == 'GAP':
                month_field_map[current_month]['GAP'] = idx
    
    return month_field_map

def process_single_month(template_path, month_name, data_rows, columns, file_info, output_path, month_field_map=None):
    """Process a single month and save to output path."""
    try:
        doc = Document(template_path)
        month_num = get_month_number(month_name)
        term_month_index = get_term_month_index(month_name, term=file_info['term'])
        replacements = {
            "{{year}}": file_info['year_range'],
            "{{act_mon}}": month_num,
            "{{term_mon}}": term_month_index,
            "{{month}}": month_name,
            "ES/00": f"ES/{file_info['standard']}"
        }
        for paragraph in doc.paragraphs:
            replace_placeholders_in_paragraph(paragraph, replacements)
        target_table = None
        placeholder_row_idx = None
        col_map = None
        for table in doc.tables:
            col_map, row_idx = get_placeholder_col_map_and_row(table)
            if col_map:
                target_table = table
                placeholder_row_idx = row_idx
                break
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholders_in_paragraph(paragraph, replacements)
        if not target_table or not col_map:
            logger.error("Could not find target table or placeholder row in template")
            return False
        target_table._tbl.remove(target_table.rows[placeholder_row_idx]._tr)
        is_syjc = file_info['original_std'] == 'SYJC'
        is_fyjc = file_info['original_std'] == 'FYJC'
        first_data_row = placeholder_row_idx
        # Use the correct field_map: for TOTAL page, always use 'TOTAL' as the key
        if month_field_map:
            if month_name.upper() not in month_field_map and 'TOTAL' in month_field_map:
                field_map = month_field_map['TOTAL']
            else:
                field_map = month_field_map.get(month_name.upper(), columns)
        else:
            field_map = columns
        for data_idx, data_row in enumerate(data_rows):
            if not data_row or len(data_row) < 2:
                continue
            if data_row[0].strip().upper() == 'TOTAL':
                break
            table_row_idx = first_data_row + data_idx
            if table_row_idx >= len(target_table.rows):
                logger.warning(f"Not enough rows in table for data row {data_idx+1}, adding row.")
                target_table.add_row()
            row = target_table.rows[table_row_idx]
            # Fill SR.NO. and INITIALS
            if col_map.get('{{col_srno}}') is not None:
                row.cells[col_map['{{col_srno}}']].text = data_row[0].strip()
            if col_map.get('{{col_initials}}') is not None and len(data_row) > 1:
                row.cells[col_map['{{col_initials}}']].text = data_row[1].strip()
            # Find indices for ALLOTTED, ENGAGED, GAP in the CSV for this month
            allotted_idx = field_map.get('ALLOTTED')
            engaged_idx = field_map.get('ENGAGED')
            gap_idx = field_map.get('GAP')
            # Fill XI columns (FYJC or always if you want to show --)
            if is_fyjc or not is_syjc:
                if col_map.get('{{col_xi_allotted}}') is not None and allotted_idx is not None and len(data_row) > allotted_idx:
                    row.cells[col_map['{{col_xi_allotted}}']].text = data_row[allotted_idx]
                else:
                    row.cells[col_map['{{col_xi_allotted}}']].text = '--'
                if col_map.get('{{col_xi_engaged}}') is not None and engaged_idx is not None and len(data_row) > engaged_idx:
                    row.cells[col_map['{{col_xi_engaged}}']].text = data_row[engaged_idx]
                else:
                    row.cells[col_map['{{col_xi_engaged}}']].text = '--'
                if col_map.get('{{col_xi_gap}}') is not None and gap_idx is not None and len(data_row) > gap_idx:
                    val = data_row[gap_idx]
                    row.cells[col_map['{{col_xi_gap}}']].text = val
                else:
                    row.cells[col_map['{{col_xi_gap}}']].text = '--'
            else:
                if col_map.get('{{col_xi_allotted}}') is not None:
                    row.cells[col_map['{{col_xi_allotted}}']].text = '--'
                if col_map.get('{{col_xi_engaged}}') is not None:
                    row.cells[col_map['{{col_xi_engaged}}']].text = '--'
                if col_map.get('{{col_xi_gap}}') is not None:
                    row.cells[col_map['{{col_xi_gap}}']].text = '--'
            # Fill XII columns (SYJC or always if you want to show --)
            if is_syjc or not is_fyjc:
                if col_map.get('{{col_xii_allotted}}') is not None and allotted_idx is not None and len(data_row) > allotted_idx:
                    row.cells[col_map['{{col_xii_allotted}}']].text = data_row[allotted_idx]
                else:
                    row.cells[col_map['{{col_xii_allotted}}']].text = '--'
                if col_map.get('{{col_xii_engaged}}') is not None and engaged_idx is not None and len(data_row) > engaged_idx:
                    row.cells[col_map['{{col_xii_engaged}}']].text = data_row[engaged_idx]
                else:
                    row.cells[col_map['{{col_xii_engaged}}']].text = '--'
                if col_map.get('{{col_xii_gap}}') is not None and gap_idx is not None and len(data_row) > gap_idx:
                    val = data_row[gap_idx]
                    row.cells[col_map['{{col_xii_gap}}']].text = val
                else:
                    row.cells[col_map['{{col_xii_gap}}']].text = '--'
            else:
                if col_map.get('{{col_xii_allotted}}') is not None:
                    row.cells[col_map['{{col_xii_allotted}}']].text = '--'
                if col_map.get('{{col_xii_engaged}}') is not None:
                    row.cells[col_map['{{col_xii_engaged}}']].text = '--'
                if col_map.get('{{col_xii_gap}}') is not None:
                    row.cells[col_map['{{col_xii_gap}}']].text = '--'
        doc.save(output_path)
        logger.info(f"Processed {month_name} and saved to {output_path}")
        return True
    except Exception as e:
        logger.error(f"Error processing {month_name}: {e}")
        return False

def merge_with_win32com(docx_files, output_file):
    """Merge docx files using Microsoft Word automation (win32com)."""
    if not WIN32COM_AVAILABLE:
        logger.error("win32com is not available. Please install pywin32.")
        return False
    if not docx_files:
        logger.error("No files to merge.")
        return False
    logger.info("Merging files with Microsoft Word automation...")
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    try:
        # Open the first document
        main_doc = word.Documents.Open(os.path.abspath(docx_files[0]))
        for doc_path in docx_files[1:]:
            main_doc.Activate()
            # Insert a page break, then insert the next file
            word.Selection.EndKey(Unit=6)  # Move to end
            word.Selection.InsertBreak(7)  # wdPageBreak
            word.Selection.InsertFile(os.path.abspath(doc_path))
        # Save as output_file
        main_doc.SaveAs(os.path.abspath(output_file))
        main_doc.Close()
        logger.info(f"Merged file saved as: {output_file}")
        return True
    except Exception as e:
        logger.error(f"Error during win32com merge: {e}")
        return False
    finally:
        word.Quit()

def create_multi_month_document(csv_path, template_path, output_folder="output_word_files"):
    """Create individual month files and then provide instructions to manually combine them."""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    # Create month-specific folder for individual month files
    month_folder = os.path.join(output_folder, "month_files")
    if not os.path.exists(month_folder):
        os.makedirs(month_folder)
    
    file_info = parse_filename(csv_path)
    if not file_info:
        logger.error(f"Could not parse information from filename: {csv_path}")
        return False
    
    logger.info(f"Processing file: {csv_path}")
    logger.info(f"File info: {file_info}")
    
    month_files = []
    
    try:
        # Open with UTF-8-SIG to handle BOM
        with open(csv_path, 'r', newline='', encoding='utf-8-sig') as csvfile:
            reader = csv.reader(csvfile)
            
            # First row contains months
            header_row = next(reader)
            
            # Second row contains column labels
            field_row = next(reader)
            
            # Map column indices for each month
            month_field_map = parse_month_field_columns(header_row, field_row)
            
            logger.info(f"Month field mapping: {month_field_map}")
            
            # Store all data rows
            all_data = []
            for row in reader:
                if row and len(row) > 0 and row[0].strip():
                    all_data.append(row)
                    if row[0].strip().upper() == 'TOTAL':
                        break
            
            # In create_multi_month_document, after reading header_row and field_row, determine the first and last month dynamically
            months_in_file = [cell.strip().upper() for cell in header_row[2:] if cell.strip() and cell.strip().upper() not in ("TOTAL", "SR.NO.", "INITIALS")]
            if months_in_file:
                start_month = months_in_file[0]
                end_month = months_in_file[-1]
                month_range = f"{start_month}-{end_month}"
            else:
                month_range = ""
            
            # Process each month
            for month_name, columns in month_field_map.items():
                if not all(key in columns for key in ['ALLOTTED', 'ENGAGED', 'GAP']):
                    logger.warning(f"Skipping month {month_name} - missing columns: {columns}")
                    continue
                # If this is the TOTAL page, use 'TOTAL' for lookups but pass month_range for the Word placeholder
                if month_name == 'TOTAL':
                    total_file = os.path.join(month_folder, f"{file_info['original_std']}_TOTAL_{file_info['year_range']}.docx")
                    success = process_single_month(
                        template_path=template_path,
                        month_name=month_range,  # For Word placeholder
                        data_rows=all_data,
                        columns=month_field_map['TOTAL'],  # For data/column lookups
                        file_info=file_info,
                        output_path=total_file,
                        month_field_map=month_field_map
                    )
                    if success:
                        month_files.append(total_file)
                    continue
                # Process the month as usual
                month_file = os.path.join(month_folder, f"{file_info['original_std']}_{month_name}_{file_info['year_range']}.docx")
                success = process_single_month(
                    template_path=template_path,
                    month_name=month_name,
                    data_rows=all_data,
                    columns=columns,
                    file_info=file_info,
                    output_path=month_file,
                    month_field_map=month_field_map
                )
                if success:
                    month_files.append(month_file)
            
            # Create a text file with instructions on how to manually combine the files
            if month_files:
                combined_filename = f"{file_info['original_std']}_{file_info['year_range']}_term{file_info['term']}_all_months.docx"
                combined_path = os.path.join(output_folder, combined_filename)
                if WIN32COM_AVAILABLE:
                    merge_with_win32com(month_files, combined_path)
                else:
                    instruction_file = os.path.join(output_folder, f"{combined_filename}_INSTRUCTIONS.txt")
                    with open(instruction_file, 'w') as f:
                        f.write("INSTRUCTIONS FOR COMBINING THE MONTH FILES\n")
                        f.write("===========================================\n\n")
                        f.write("Due to formatting issues when programmatically combining Word documents,\n")
                        f.write("please follow these steps to combine the month files manually:\n\n")
                        f.write("1. Open the first month file\n")
                        f.write("2. For each additional month file:\n")
                        f.write("   a. Place cursor at the end of the document\n")
                        f.write("   b. Insert -> Page Break\n")
                        f.write("   c. Insert -> Object -> Text from file\n")
                        f.write("   d. Select the next month file\n")
                        f.write("3. Save the combined document\n\n")
                        f.write("Individual month files are located in the 'month_files' folder:\n\n")
                        
                        for i, file_path in enumerate(month_files, 1):
                            f.write(f"{i}. {os.path.basename(file_path)}\n")
                
                logger.info(f"Created {len(month_files)} individual month files in the '{month_folder}' folder")
                logger.info(f"Instructions for manually combining files written to {instruction_file}")
            else:
                logger.error("No month files were created")
    
    except Exception as e:
        logger.error(f"Error processing file {csv_path}: {e}")
        return False
    
    return True

def process_excel_files(excel_folder="excel_copies", template_path="executive_summary_template.docx", output_folder="output_word_files"):
    """Process all Excel files in the given folder."""
    if not os.path.exists(excel_folder):
        logger.error(f"Excel folder not found: {excel_folder}")
        return
    
    if not os.path.exists(template_path):
        logger.error(f"Template file not found: {template_path}")
        return
    
    for filename in os.listdir(excel_folder):
        if filename.endswith('.csv') and filename.startswith('iso_excel_'):
            file_path = os.path.join(excel_folder, filename)
            logger.info(f"Processing: {filename}")
            create_multi_month_document(file_path, template_path, output_folder)

# Example usage
if __name__ == "__main__":
    logger.info("Starting Excel to Word conversion")
    process_excel_files()
    logger.info("Conversion completed")
