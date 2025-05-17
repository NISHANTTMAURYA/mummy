import os
import csv
import re
import logging
from docx import Document
import shutil
from copy import deepcopy
import pandas as pd
import tempfile

# Add win32com import for Word automation
try:
    import win32com.client
    WIN32COM_AVAILABLE = True
except ImportError:
    WIN32COM_AVAILABLE = False

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def parse_filename(filename):
    """Extract information from the filename."""
    # Expected format: iso_excel_YYYY-YYYY_termN_STD.csv
    pattern = r'iso_excel_(\d{4}-\d{4})_term(\d+)_(\w+)\.(csv|xlsx|xls)'
    match = re.match(pattern, os.path.basename(filename))
    
    if match:
        year_range, term, std, _ = match.groups()
        # Convert STD to actual standard (FYJC=XI and SYJC=XII)
        std_mapping = {"FYJC": "XI", "SYJC": "XII"}
        standard = std_mapping.get(std, std)
        
        return {
            'year_range': year_range,
            'term': term,
            'standard': standard,
            'original_std': std
        }
    return None

def get_month_number(month_name):
    """Convert month name to number."""
    month_mapping = {
        'JUNE': '06', 'JULY': '07', 'AUG': '08', 'SEPTEMBER': '09', 
        'SEP': '09', 'OCTOBER': '10', 'OCT': '10', 'NOVEMBER': '11',
        'NOV': '11', 'DECEMBER': '12', 'DEC': '12', 'JANUARY': '01',
        'JAN': '01', 'FEBRUARY': '02', 'FEB': '02', 'MARCH': '03',
        'MAR': '03', 'APRIL': '04', 'APR': '04', 'MAY': '05'
    }
    return month_mapping.get(month_name.upper(), '00')

def get_term_month_index(month_name, term):
    """Get the month index within the term."""
    if term == '1':
        months = ['JUNE', 'JULY', 'AUG', 'SEP', 'OCT']
    else:  # term 2
        months = ['NOV', 'DEC', 'JAN', 'FEB']
    
    for i, m in enumerate(months, 1):
        if month_name.upper().startswith(m):
            return f"{i:02d}"
    return "01"  # Default

def replace_placeholders_in_paragraph(paragraph, replacements):
    """Replace placeholders in paragraph text without affecting other formatting."""
    if not any(placeholder in paragraph.text for placeholder in replacements.keys()):
        return False
    
    # Create a new paragraph text with replacements
    text = paragraph.text
    for placeholder, value in replacements.items():
        if placeholder in text:
            text = text.replace(placeholder, value)
    
    # Replace text in the paragraph
    paragraph.text = text
    return True

def find_standard_columns(table):
    """Find the column indices for XI and XII (ALLOTTED, ENGAGED, GAP) by header text."""
    xi_cols = {'ALLOTTED': None, 'ENGAGED': None, 'GAP': None}
    xii_cols = {'ALLOTTED': None, 'ENGAGED': None, 'GAP': None}
    xi_start, xii_start = None, None
    for row in table.rows[:3]:
        for idx, cell in enumerate(row.cells):
            txt = cell.text.strip().upper()
            if txt == 'XI':
                xi_start = idx
            elif txt == 'XII':
                xii_start = idx
    for row in table.rows[:4]:
        for idx, cell in enumerate(row.cells):
            txt = cell.text.strip().upper()
            if xi_start is not None and xi_start <= idx < xi_start + 3:
                if txt == 'ALLOTTED':
                    xi_cols['ALLOTTED'] = idx
                elif txt == 'ENGAGED':
                    xi_cols['ENGAGED'] = idx
                elif txt == 'GAP':
                    xi_cols['GAP'] = idx
            if xii_start is not None and xii_start <= idx < xii_start + 3:
                if txt == 'ALLOTTED':
                    xii_cols['ALLOTTED'] = idx
                elif txt == 'ENGAGED':
                    xii_cols['ENGAGED'] = idx
                elif txt == 'GAP':
                    xii_cols['GAP'] = idx
    logger.info(f"XI columns: {xi_cols}, XII columns: {xii_cols}")
    return xi_cols, xii_cols

def get_placeholder_col_map_and_row(table):
    """Return (placeholder_map, row_index) for the row with all placeholders."""
    for row_idx, row in enumerate(table.rows):
        placeholders = {}
        all_placeholders = True
        for idx, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt.startswith('{{col_') and txt.endswith('}}'):
                placeholders[txt] = idx
            else:
                all_placeholders = False
        if all_placeholders and placeholders:
            logger.info(f"Found placeholder row at index {row_idx}: {placeholders}")
            return placeholders, row_idx
    logger.error("No placeholder row found in table!")
    return None, None

def parse_month_field_columns(header_row, field_row):
    """Return a dict: {month: {field: col_idx}} for each month and field in the CSV."""
    month_field_map = {}
    current_month = None
    month_start_idx = None
    
    # Skip the first two columns (SR.NO. and INITIALS)
    for idx, (month_cell, field_cell) in enumerate(zip(header_row[2:], field_row[2:]), start=2):
        # If the month cell is not empty, update current_month
        if month_cell.strip():
            current_month = month_cell.strip().upper()
            if current_month not in month_field_map:
                month_field_map[current_month] = {}
                month_start_idx = idx
        
        # Only map if we have a current_month and a valid field
        if current_month and month_start_idx is not None:
            field = field_cell.strip().upper()
            if field == 'ALOTTED':
                month_field_map[current_month]['ALLOTTED'] = idx
            elif field == 'ENGAGED':
                month_field_map[current_month]['ENGAGED'] = idx
            elif field == 'GAP':
                month_field_map[current_month]['GAP'] = idx
    
    return month_field_map

def process_single_month(template_path, month_name, data_rows1, data_rows2, columns1, columns2, file_info1, file_info2, output_path):
    """Process a single month and save to output path."""
    try:
        doc = Document(template_path)
        month_num = get_month_number(month_name)
        term_month_index = get_term_month_index(month_name, term=file_info1['term'])
        
        # For TOTAL page, format the month name based on term
        if month_name.upper() == 'TOTAL':
            if file_info1['term'] == '1':
                month_name = "JUNE-OCT"
            else:  # term 2
                month_name = "NOV-FEB"
        
        replacements = {
            "{{year}}": file_info1['year_range'],
            "{{act_mon}}": month_num,
            "{{term_mon}}": term_month_index,
            "{{month}}": month_name,
            "ES/00": f"ES/{file_info1['standard']}"
        }
        
        for paragraph in doc.paragraphs:
            replace_placeholders_in_paragraph(paragraph, replacements)
        target_table = None
        placeholder_row_idx = None
        col_map = None
        for table in doc.tables:
            col_map, row_idx = get_placeholder_col_map_and_row(table)
            if col_map:
                target_table = table
                placeholder_row_idx = row_idx
                break
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholders_in_paragraph(paragraph, replacements)
        if not target_table or not col_map:
            logger.error("Could not find target table or placeholder row in template")
            return False
        target_table._tbl.remove(target_table.rows[placeholder_row_idx]._tr)
        first_data_row = placeholder_row_idx

        # Get field maps for both files
        if month_name.upper() not in columns1 and 'TOTAL' in columns1:
            field_map1 = columns1['TOTAL']
        else:
            field_map1 = columns1.get(month_name.upper(), columns1)
        
        if month_name.upper() not in columns2 and 'TOTAL' in columns2:
            field_map2 = columns2['TOTAL']
        else:
            field_map2 = columns2.get(month_name.upper(), columns2)

        # Check if this is a single file process (same file info)
        is_single_file = file_info1 == file_info2

        # Process data from both files
        for data_idx, (data_row1, data_row2) in enumerate(zip(data_rows1, data_rows2)):
            if not data_row1 or len(data_row1) < 2 or not data_row2 or len(data_row2) < 2:
                continue
            # For TOTAL page, look for the row that starts with "TOTAL"
            if month_name.upper() == 'TOTAL' and data_row1[0].strip().upper() != 'TOTAL':
                continue
            # For regular months, stop at TOTAL row
            elif month_name.upper() != 'TOTAL' and data_row1[0].strip().upper() == 'TOTAL':
                break

            table_row_idx = first_data_row + data_idx
            if table_row_idx >= len(target_table.rows):
                logger.warning(f"Not enough rows in table for data row {data_idx+1}, adding row.")
                target_table.add_row()
            
            row = target_table.rows[table_row_idx]

            # Fill SR.NO. and INITIALS (use data from first file)
            if col_map.get('{{col_srno}}') is not None:
                row.cells[col_map['{{col_srno}}']].text = data_row1[0].strip()
            if col_map.get('{{col_initials}}') is not None and len(data_row1) > 1:
                row.cells[col_map['{{col_initials}}']].text = data_row1[1].strip()

            # Get indices for ALLOTTED, ENGAGED, GAP for both files
            allotted_idx1 = field_map1.get('ALLOTTED')
            engaged_idx1 = field_map1.get('ENGAGED')
            gap_idx1 = field_map1.get('GAP')
            
            allotted_idx2 = field_map2.get('ALLOTTED')
            engaged_idx2 = field_map2.get('ENGAGED')
            gap_idx2 = field_map2.get('GAP')

            # For single file processing, only fill the appropriate columns based on standard
            if is_single_file:
                if file_info1['original_std'] == 'FYJC':
                    # Fill XI columns only
                    if col_map.get('{{col_xi_allotted}}') is not None and allotted_idx1 is not None and len(data_row1) > allotted_idx1:
                        row.cells[col_map['{{col_xi_allotted}}']].text = data_row1[allotted_idx1]
                    else:
                        row.cells[col_map['{{col_xi_allotted}}']].text = '--'

                    if col_map.get('{{col_xi_engaged}}') is not None and engaged_idx1 is not None and len(data_row1) > engaged_idx1:
                        row.cells[col_map['{{col_xi_engaged}}']].text = data_row1[engaged_idx1]
                    else:
                        row.cells[col_map['{{col_xi_engaged}}']].text = '--'

                    if col_map.get('{{col_xi_gap}}') is not None and gap_idx1 is not None and len(data_row1) > gap_idx1:
                        row.cells[col_map['{{col_xi_gap}}']].text = data_row1[gap_idx1]
                    else:
                        row.cells[col_map['{{col_xi_gap}}']].text = '--'

                    # Clear XII columns
                    if col_map.get('{{col_xii_allotted}}') is not None:
                        row.cells[col_map['{{col_xii_allotted}}']].text = '--'
                    if col_map.get('{{col_xii_engaged}}') is not None:
                        row.cells[col_map['{{col_xii_engaged}}']].text = '--'
                    if col_map.get('{{col_xii_gap}}') is not None:
                        row.cells[col_map['{{col_xii_gap}}']].text = '--'
                else:  # SYJC
                    # Fill XII columns only
                    if col_map.get('{{col_xii_allotted}}') is not None and allotted_idx1 is not None and len(data_row1) > allotted_idx1:
                        row.cells[col_map['{{col_xii_allotted}}']].text = data_row1[allotted_idx1]
                    else:
                        row.cells[col_map['{{col_xii_allotted}}']].text = '--'

                    if col_map.get('{{col_xii_engaged}}') is not None and engaged_idx1 is not None and len(data_row1) > engaged_idx1:
                        row.cells[col_map['{{col_xii_engaged}}']].text = data_row1[engaged_idx1]
                    else:
                        row.cells[col_map['{{col_xii_engaged}}']].text = '--'

                    if col_map.get('{{col_xii_gap}}') is not None and gap_idx1 is not None and len(data_row1) > gap_idx1:
                        row.cells[col_map['{{col_xii_gap}}']].text = data_row1[gap_idx1]
                    else:
                        row.cells[col_map['{{col_xii_gap}}']].text = '--'

                    # Clear XI columns
                    if col_map.get('{{col_xi_allotted}}') is not None:
                        row.cells[col_map['{{col_xi_allotted}}']].text = '--'
                    if col_map.get('{{col_xi_engaged}}') is not None:
                        row.cells[col_map['{{col_xi_engaged}}']].text = '--'
                    if col_map.get('{{col_xi_gap}}') is not None:
                        row.cells[col_map['{{col_xi_gap}}']].text = '--'
            else:
                # Original dual file processing logic
                # Fill XI columns (FYJC data)
                if col_map.get('{{col_xi_allotted}}') is not None and allotted_idx1 is not None and len(data_row1) > allotted_idx1:
                    row.cells[col_map['{{col_xi_allotted}}']].text = data_row1[allotted_idx1]
                else:
                    row.cells[col_map['{{col_xi_allotted}}']].text = '--'

                if col_map.get('{{col_xi_engaged}}') is not None and engaged_idx1 is not None and len(data_row1) > engaged_idx1:
                    row.cells[col_map['{{col_xi_engaged}}']].text = data_row1[engaged_idx1]
                else:
                    row.cells[col_map['{{col_xi_engaged}}']].text = '--'

                if col_map.get('{{col_xi_gap}}') is not None and gap_idx1 is not None and len(data_row1) > gap_idx1:
                    row.cells[col_map['{{col_xi_gap}}']].text = data_row1[gap_idx1]
                else:
                    row.cells[col_map['{{col_xi_gap}}']].text = '--'

                # Fill XII columns (SYJC data)
                if col_map.get('{{col_xii_allotted}}') is not None and allotted_idx2 is not None and len(data_row2) > allotted_idx2:
                    row.cells[col_map['{{col_xii_allotted}}']].text = data_row2[allotted_idx2]
                else:
                    row.cells[col_map['{{col_xii_allotted}}']].text = '--'

                if col_map.get('{{col_xii_engaged}}') is not None and engaged_idx2 is not None and len(data_row2) > engaged_idx2:
                    row.cells[col_map['{{col_xii_engaged}}']].text = data_row2[engaged_idx2]
                else:
                    row.cells[col_map['{{col_xii_engaged}}']].text = '--'

                if col_map.get('{{col_xii_gap}}') is not None and gap_idx2 is not None and len(data_row2) > gap_idx2:
                    row.cells[col_map['{{col_xii_gap}}']].text = data_row2[gap_idx2]
                else:
                    row.cells[col_map['{{col_xii_gap}}']].text = '--'

        # --- NEW: Fill the TOTAL row at the end of the table for each month ---
        # Only for non-TOTAL pages (since TOTAL page already handled)
        if month_name.upper() != 'TOTAL':
            # Find the TOTAL row in the data
            total_row1 = next((row for row in data_rows1 if row and row[0].strip().upper() == 'TOTAL'), None)
            total_row2 = next((row for row in data_rows2 if row and row[0].strip().upper() == 'TOTAL'), None)
            # Find the last row in the table (should be labeled TOTAL)
            last_row = target_table.rows[-1]
            # Fill SR.NO. and INITIALS as 'TOTAL' and blank
            if col_map.get('{{col_srno}}') is not None:
                last_row.cells[col_map['{{col_srno}}']].text = 'TOTAL'
            if col_map.get('{{col_initials}}') is not None:
                last_row.cells[col_map['{{col_initials}}']].text = ''
            # Fill XI and XII columns appropriately for single or dual file
            if is_single_file:
                if file_info1['original_std'] == 'FYJC':
                    # Fill XI columns from total_row1
                    if total_row1:
                        allotted_idx1 = field_map1.get('ALLOTTED')
                        engaged_idx1 = field_map1.get('ENGAGED')
                        gap_idx1 = field_map1.get('GAP')
                        if col_map.get('{{col_xi_allotted}}') is not None and allotted_idx1 is not None and len(total_row1) > allotted_idx1:
                            last_row.cells[col_map['{{col_xi_allotted}}']].text = total_row1[allotted_idx1]
                        if col_map.get('{{col_xi_engaged}}') is not None and engaged_idx1 is not None and len(total_row1) > engaged_idx1:
                            last_row.cells[col_map['{{col_xi_engaged}}']].text = total_row1[engaged_idx1]
                        if col_map.get('{{col_xi_gap}}') is not None and gap_idx1 is not None and len(total_row1) > gap_idx1:
                            last_row.cells[col_map['{{col_xi_gap}}']].text = total_row1[gap_idx1]
                    # Set XII columns to '--'
                    if col_map.get('{{col_xii_allotted}}') is not None:
                        last_row.cells[col_map['{{col_xii_allotted}}']].text = '--'
                    if col_map.get('{{col_xii_engaged}}') is not None:
                        last_row.cells[col_map['{{col_xii_engaged}}']].text = '--'
                    if col_map.get('{{col_xii_gap}}') is not None:
                        last_row.cells[col_map['{{col_xii_gap}}']].text = '--'
                else:  # SYJC
                    # Fill XII columns from total_row1
                    if total_row1:
                        allotted_idx1 = field_map1.get('ALLOTTED')
                        engaged_idx1 = field_map1.get('ENGAGED')
                        gap_idx1 = field_map1.get('GAP')
                        if col_map.get('{{col_xii_allotted}}') is not None and allotted_idx1 is not None and len(total_row1) > allotted_idx1:
                            last_row.cells[col_map['{{col_xii_allotted}}']].text = total_row1[allotted_idx1]
                        if col_map.get('{{col_xii_engaged}}') is not None and engaged_idx1 is not None and len(total_row1) > engaged_idx1:
                            last_row.cells[col_map['{{col_xii_engaged}}']].text = total_row1[engaged_idx1]
                        if col_map.get('{{col_xii_gap}}') is not None and gap_idx1 is not None and len(total_row1) > gap_idx1:
                            last_row.cells[col_map['{{col_xii_gap}}']].text = total_row1[gap_idx1]
                    # Set XI columns to '--'
                    if col_map.get('{{col_xi_allotted}}') is not None:
                        last_row.cells[col_map['{{col_xi_allotted}}']].text = '--'
                    if col_map.get('{{col_xi_engaged}}') is not None:
                        last_row.cells[col_map['{{col_xi_engaged}}']].text = '--'
                    if col_map.get('{{col_xi_gap}}') is not None:
                        last_row.cells[col_map['{{col_xi_gap}}']].text = '--'
            else:
                # Dual file logic (unchanged)
                if total_row1:
                    allotted_idx1 = field_map1.get('ALLOTTED')
                    engaged_idx1 = field_map1.get('ENGAGED')
                    gap_idx1 = field_map1.get('GAP')
                    if col_map.get('{{col_xi_allotted}}') is not None and allotted_idx1 is not None and len(total_row1) > allotted_idx1:
                        last_row.cells[col_map['{{col_xi_allotted}}']].text = total_row1[allotted_idx1]
                    if col_map.get('{{col_xi_engaged}}') is not None and engaged_idx1 is not None and len(total_row1) > engaged_idx1:
                        last_row.cells[col_map['{{col_xi_engaged}}']].text = total_row1[engaged_idx1]
                    if col_map.get('{{col_xi_gap}}') is not None and gap_idx1 is not None and len(total_row1) > gap_idx1:
                        last_row.cells[col_map['{{col_xi_gap}}']].text = total_row1[gap_idx1]
                if total_row2:
                    allotted_idx2 = field_map2.get('ALLOTTED')
                    engaged_idx2 = field_map2.get('ENGAGED')
                    gap_idx2 = field_map2.get('GAP')
                    if col_map.get('{{col_xii_allotted}}') is not None and allotted_idx2 is not None and len(total_row2) > allotted_idx2:
                        last_row.cells[col_map['{{col_xii_allotted}}']].text = total_row2[allotted_idx2]
                    if col_map.get('{{col_xii_engaged}}') is not None and engaged_idx2 is not None and len(total_row2) > engaged_idx2:
                        last_row.cells[col_map['{{col_xii_engaged}}']].text = total_row2[engaged_idx2]
                    if col_map.get('{{col_xii_gap}}') is not None and gap_idx2 is not None and len(total_row2) > gap_idx2:
                        last_row.cells[col_map['{{col_xii_gap}}']].text = total_row2[gap_idx2]
        # --- END NEW ---
        doc.save(output_path)
        logger.info(f"Processed {month_name} and saved to {output_path}")
        return True
    except Exception as e:
        logger.error(f"Error processing {month_name}: {e}")
        return False

def merge_with_win32com(docx_files, output_file):
    """Merge docx files using Microsoft Word automation (win32com)."""
    if not WIN32COM_AVAILABLE:
        logger.error("win32com is not available. Please install pywin32.")
        return False
    if not docx_files:
        logger.error("No files to merge.")
        return False
    logger.info("Merging files with Microsoft Word automation...")
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    try:
        # Open the first document
        main_doc = word.Documents.Open(os.path.abspath(docx_files[0]))
        for doc_path in docx_files[1:]:
            main_doc.Activate()
            # Insert a page break, then insert the next file
            word.Selection.EndKey(Unit=6)  # Move to end
            word.Selection.InsertBreak(7)  # wdPageBreak
            word.Selection.InsertFile(os.path.abspath(doc_path))
        # Save as output_file
        main_doc.SaveAs(os.path.abspath(output_file))
        main_doc.Close()
        logger.info(f"Merged file saved as: {output_file}")
        return True
    except Exception as e:
        logger.error(f"Error during win32com merge: {e}")
        return False
    finally:
        word.Quit()

def create_multi_month_document(csv_path1, csv_path2, template_path, output_folder="output_word_files"):
    """Create individual month files and then provide instructions to manually combine them."""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    # Create month-specific folder for individual month files
    month_folder = os.path.join(output_folder, "month_files")
    if not os.path.exists(month_folder):
        os.makedirs(month_folder)
    
    file1_info = parse_filename(csv_path1)
    file2_info = parse_filename(csv_path2)
    
    if not file1_info or not file2_info:
        logger.error(f"Could not parse information from filenames: {csv_path1} or {csv_path2}")
        return False
    
    logger.info(f"Processing files: {csv_path1} and {csv_path2}")
    logger.info(f"File 1 info: {file1_info}")
    logger.info(f"File 2 info: {file2_info}")
    
    # Define proper month order
    month_order = ['JUNE', 'JULY', 'AUG', 'SEP', 'OCT', 'NOV', 'DEC', 'JAN', 'FEB', 'MAR', 'APR', 'MAY']
    month_files = []
    
    try:
        # Open with UTF-8-SIG to handle BOM
        with open(csv_path1, 'r', newline='', encoding='utf-8-sig') as csvfile1, open(csv_path2, 'r', newline='', encoding='utf-8-sig') as csvfile2:
            reader1 = csv.reader(csvfile1)
            reader2 = csv.reader(csvfile2)
            
            # First row contains months
            header_row1 = next(reader1)
            header_row2 = next(reader2)
            
            # Second row contains column labels
            field_row1 = next(reader1)
            field_row2 = next(reader2)
            
            # Map column indices for each month
            month_field_map1 = parse_month_field_columns(header_row1, field_row1)
            month_field_map2 = parse_month_field_columns(header_row2, field_row2)
            
            logger.info(f"Month field mapping for file 1: {month_field_map1}")
            logger.info(f"Month field mapping for file 2: {month_field_map2}")
            
            # Store all data rows
            all_data1 = []
            all_data2 = []
            for row in reader1:
                if row and len(row) > 0 and row[0].strip():
                    all_data1.append(row)
                    if row[0].strip().upper() == 'TOTAL':
                        break
            
            for row in reader2:
                if row and len(row) > 0 and row[0].strip():
                    all_data2.append(row)
                    if row[0].strip().upper() == 'TOTAL':
                        break
            
            # Get common months from both files and sort them according to month_order
            common_months = sorted(
                set(month_field_map1.keys()) & set(month_field_map2.keys()),
                key=lambda x: month_order.index(x) if x in month_order else float('inf')
            )
            
            # Process each month in the correct order
            for month_name in common_months:
                if month_name in month_field_map1 and month_name in month_field_map2:
                    # Process data from both files for this month
                    month_file = os.path.join(month_folder, f"{month_name}_{file1_info['year_range']}.docx")
                    success = process_single_month(
                        template_path=template_path,
                        month_name=month_name,
                        data_rows1=all_data1,
                        data_rows2=all_data2,
                        columns1=month_field_map1[month_name],
                        columns2=month_field_map2[month_name],
                        file_info1=file1_info,
                        file_info2=file2_info,
                        output_path=month_file
                    )
                    if success:
                        month_files.append(month_file)
    
    except Exception as e:
        logger.error(f"Error processing files: {e}")
        return False
    
    return True

def convert_excel_to_csv(excel_path):
    """Convert Excel file to CSV file with the same base name and return the path."""
    csv_path = os.path.splitext(excel_path)[0] + ".csv"
    try:
        logger.info(f"Reading Excel file: {excel_path}")
        df = pd.read_excel(excel_path, header=None)
        logger.info(f"Converting to CSV: {csv_path}")
        df.to_csv(csv_path, index=False, header=False, encoding='utf-8-sig')
        return csv_path
    except Exception as e:
        logger.error(f"Error converting Excel to CSV: {e}")
        return None

def force_excel_recalc_and_save(excel_path):
    """Open Excel file, force recalculation, and save using Excel COM automation."""
    try:
        import win32com.client
        excel = win32com.client.Dispatch("Excel.Application")
        excel.Visible = False
        wb = excel.Workbooks.Open(os.path.abspath(excel_path))
        wb.RefreshAll()
        excel.CalculateFullRebuild()
        wb.Save()
        wb.Close(False)
        excel.Quit()
        logger.info(f"Excel formulas recalculated and file saved: {excel_path}")
    except Exception as e:
        logger.error(f"Error recalculating and saving Excel file: {e}")

def process_excel_files(excel_folder="excel_copies", template_path="executive_summary_template.docx", output_folder="output_word_files"):
    """Process all Excel files in the given folder."""
    if not os.path.exists(excel_folder):
        logger.error(f"Excel folder not found: {excel_folder}")
        return
    if not os.path.exists(template_path):
        logger.error(f"Template file not found: {template_path}")
        return
    for filename in os.listdir(excel_folder):
        if filename.endswith(('.xlsx', '.xls')) and filename.startswith('iso_excel_'):
            file_path = os.path.join(excel_folder, filename)
            logger.info(f"Processing Excel file: {filename}")
            force_excel_recalc_and_save(file_path)
            csv_path = convert_excel_to_csv(file_path)
            if csv_path:
                try:
                    create_multi_month_document(csv_path, csv_path, template_path, output_folder)
                finally:
                    try:
                        os.unlink(csv_path)
                        logger.info(f"Cleaned up CSV file: {csv_path}")
                    except Exception as e:
                        logger.error(f"Error cleaning up CSV file: {e}")

def process_single_excel_file(excel_path, template_path="executive_summary_template.docx", output_folder="output_word_files"):
    """Process a single Excel file and return the path to the generated Word file"""
    try:
        # Create output directory if it doesn't exist
        os.makedirs(output_folder, exist_ok=True)
        
        # Convert Excel to CSV
        csv_path = convert_excel_to_csv(excel_path)
        
        # Parse filename to get year, term, and standard
        file_info = parse_filename(os.path.basename(excel_path))
        
        # Create output filename
        output_filename = f"executive_summary_{file_info['year']}_{file_info['term']}_{file_info['std']}.docx"
        output_path = os.path.join(output_folder, output_filename)
        
        # Create the Word document
        create_multi_month_document(csv_path, None, template_path, output_folder)
        
        # Return the path to the generated file
        return output_path
        
    except Exception as e:
        logger.error(f"Error processing single file: {e}")
        raise

def process_dual_excel_files(excel_path1, excel_path2, template_path="executive_summary_template.docx", output_folder="output_word_files"):
    """Process two compatible Excel files and return the path to the generated Word file"""
    try:
        # Create output directory if it doesn't exist
        os.makedirs(output_folder, exist_ok=True)
        
        # Convert both Excel files to CSV
        csv_path1 = convert_excel_to_csv(excel_path1)
        csv_path2 = convert_excel_to_csv(excel_path2)
        
        # Parse filenames to get year, term, and standards
        file_info1 = parse_filename(os.path.basename(excel_path1))
        file_info2 = parse_filename(os.path.basename(excel_path2))
        
        # Create output filename
        output_filename = f"executive_summary_{file_info1['year']}_{file_info1['term']}_combined.docx"
        output_path = os.path.join(output_folder, output_filename)
        
        # Create the Word document
        create_multi_month_document(csv_path1, csv_path2, template_path, output_folder)
        
        # Return the path to the generated file
        return output_path
        
    except Exception as e:
        logger.error(f"Error processing dual files: {e}")
        raise

# Example usage for testing
if __name__ == "__main__":
    # Test paths - replace these with your actual file paths
    excel_path1 = "excel_copies/iso_excel_2023-2024_term1_FYJC.xlsx"
    excel_path2 = "excel_copies/iso_excel_2023-2024_term1_SYJC.xlsx"  # Empty path for single file processing
    
    logger.info("Starting Excel to Word conversion")
    
    # Check if either file exists and is valid
    if excel_path1.strip() and excel_path2.strip() and os.path.exists(excel_path1) and os.path.exists(excel_path2):
        # Process both files together
        process_dual_excel_files(excel_path1, excel_path2)
    elif excel_path1.strip() and os.path.exists(excel_path1):
        # Process first file only
        process_single_excel_file(excel_path1)
    elif excel_path2.strip() and os.path.exists(excel_path2):
        # Process second file only
        process_single_excel_file(excel_path2)
    else:
        logger.error("No valid Excel files found to process")
    
    logger.info("Conversion completed")



# process_single_excel_file("path/to/your/excel_file.xlsx")