import os
import csv
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_filename(filename):
    """Extract information from the filename."""
    # Expected format: iso_excel_YYYY-YYYY_termN_STD.csv
    pattern = r'iso_excel_(\d{4}-\d{4})_term(\d+)_(\w+)\.csv'
    match = re.match(pattern, os.path.basename(filename))
    
    if match:
        year_range, term, std = match.groups()
        # Convert STD to actual standard (FYJC=XI and SYJC=XII)
        std_mapping = {"FYJC": "XI", "SYJC": "XII"}
        standard = std_mapping.get(std, std)
        
        return {
            'year_range': year_range,
            'term': term,
            'standard': standard,
            'original_std': std
        }
    return None

def get_month_number(month_name):
    """Convert month name to number."""
    month_mapping = {
        'JUNE': '06', 'JULY': '07', 'AUG': '08', 'SEPTEMBER': '09', 
        'SEP': '09', 'OCTOBER': '10', 'OCT': '10', 'NOVEMBER': '11',
        'NOV': '11', 'DECEMBER': '12', 'DEC': '12', 'JANUARY': '01',
        'JAN': '01', 'FEBRUARY': '02', 'FEB': '02', 'MARCH': '03',
        'MAR': '03', 'APRIL': '04', 'APR': '04', 'MAY': '05'
    }
    return month_mapping.get(month_name.upper(), '00')

def get_month_term_index(month_name, term):
    """Get the month index within the term."""
    if term == '1':
        months = ['JUNE', 'JULY', 'AUG', 'SEP', 'OCT']
    else:  # term 2
        months = ['NOV', 'DEC', 'JAN', 'FEB']
    
    for i, m in enumerate(months, 1):
        if month_name.upper().startswith(m):
            return f"{i:02d}"
    return "01"  # Default

def create_word_document(csv_path, output_folder="output_word_files"):
    """Create Word documents from Excel data."""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    file_info = parse_filename(csv_path)
    if not file_info:
        print(f"Error: Could not parse information from filename: {csv_path}")
        return False
    
    with open(csv_path, 'r', newline='') as csvfile:
        reader = csv.reader(csvfile)
        
        # First row contains months
        header_row = next(reader)
        
        # Second row contains column labels (ALOTTED, E-Act, etc.)
        column_labels = next(reader)
        
        # Map the column indices for each month
        month_columns = {}
        month_name = None
        
        for i, label in enumerate(header_row):
            if label and label.upper() in ['JUNE', 'JULY', 'AUG', 'SEP', 'OCT', 'NOV', 'DEC', 'JAN', 'FEB', 'MAR', 'APR', 'MAY']:
                month_name = label.upper()
                month_columns[month_name] = {}
            
            if month_name and i < len(column_labels):
                if column_labels[i] == 'ALOTTED':
                    month_columns[month_name]['ALLOTTED'] = i
                elif column_labels[i] == 'E-Act':
                    month_columns[month_name]['ENGAGED'] = i
                elif column_labels[i] == 'E-Add':
                    month_columns[month_name]['GAP'] = i
        
        # Process each month
        for month_name, columns in month_columns.items():
            if not all(key in columns for key in ['ALLOTTED', 'ENGAGED', 'GAP']):
                continue  # Skip months with incomplete data
            
            # Create a new document from template
            doc = Document('executive_summary_template.docx')
            
            # Replace placeholders in the document
            month_num = get_month_number(month_name)
            term_month_index = get_month_term_index(month_name, file_info['term'])
            
            year_code = f"{file_info['year_range']}-{month_num}/{term_month_index}"
            subject_code = "ES/ECO"  # This might need to be determined from other data
            
            # Replace placeholder text in paragraphs
            for paragraph in doc.paragraphs:
                paragraph.text = paragraph.text.replace("MONTH-JUNE", f"MONTH-{month_name.title()}")
                paragraph.text = paragraph.text.replace("2024-2025-06/J – ES/ECO/01", f"{year_code} – {subject_code}/{term_month_index}")
            
            # Find the main table to populate
            if doc.tables:
                table = doc.tables[0]  # Assuming the first table is our target
                
                # Reset row index for data rows (skip header row)
                row_idx = 1
                
                # Process each data row
                csvfile.seek(0)
                next(reader)  # Skip header row
                next(reader)  # Skip column labels
                
                for data_row in reader:
                    if not data_row or len(data_row) < 2 or not data_row[0].strip():
                        continue  # Skip empty rows
                    
                    if data_row[0].strip().upper() == 'TOTAL':
                        break  # Stop at total row
                    
                    # Add a new row if needed
                    if row_idx >= len(table.rows):
                        table.add_row()
                    
                    # Get the row in the table
                    table_row = table.rows[row_idx]
                    
                    # Fill in data
                    if len(table_row.cells) >= 4:  # Ensure the row has enough cells
                        # Serial Number
                        if len(data_row) > 0 and data_row[0]:
                            table_row.cells[0].text = data_row[0]
                        
                        # Initials
                        if len(data_row) > 1 and data_row[1]:
                            table_row.cells[1].text = data_row[1]
                        
                        # Allotted
                        if 'ALLOTTED' in columns and len(data_row) > columns['ALLOTTED']:
                            table_row.cells[2].text = data_row[columns['ALLOTTED']]
                        
                        # Engaged
                        if 'ENGAGED' in columns and len(data_row) > columns['ENGAGED']:
                            table_row.cells[3].text = data_row[columns['ENGAGED']]
                        
                        # Gap
                        if 'GAP' in columns and len(data_row) > columns['GAP'] and len(table_row.cells) > 4:
                            table_row.cells[4].text = data_row[columns['GAP']]
                    
                    row_idx += 1
            
            # Save the document
            output_filename = f"{file_info['original_std']}_{month_name}_{file_info['year_range']}.docx"
            output_path = os.path.join(output_folder, output_filename)
            doc.save(output_path)
            print(f"Created Word document: {output_path}")
    
    return True

def process_excel_files(folder_path="excel_copies", output_folder="output_word_files"):
    """Process all Excel files in the given folder."""
    if not os.path.exists(folder_path):
        print(f"Error: Folder not found: {folder_path}")
        return
    
    for filename in os.listdir(folder_path):
        if filename.endswith('.csv') and filename.startswith('iso_excel_'):
            file_path = os.path.join(folder_path, filename)
            print(f"Processing: {filename}")
            create_word_document(file_path, output_folder)

# Example usage
if __name__ == "__main__":
    # Process a single file
    # create_word_document("excel_copies/iso_excel_2023-2024_term1_SYJC.csv")
    
    # Process all Excel files in the folder
    process_excel_files()
