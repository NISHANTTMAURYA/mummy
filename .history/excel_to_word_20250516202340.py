import os
import csv
import re
import logging
from docx import Document
import shutil

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def parse_filename(filename):
    """Extract information from the filename."""
    # Expected format: iso_excel_YYYY-YYYY_termN_STD.csv
    pattern = r'iso_excel_(\d{4}-\d{4})_term(\d+)_(\w+)\.csv'
    match = re.match(pattern, os.path.basename(filename))
    
    if match:
        year_range, term, std = match.groups()
        # Convert STD to actual standard (FYJC=XI and SYJC=XII)
        std_mapping = {"FYJC": "XI", "SYJC": "XII"}
        standard = std_mapping.get(std, std)
        
        return {
            'year_range': year_range,
            'term': term,
            'standard': standard,
            'original_std': std
        }
    return None

def get_month_number(month_name):
    """Convert month name to number."""
    month_mapping = {
        'JUNE': '06', 'JULY': '07', 'AUG': '08', 'SEPTEMBER': '09', 
        'SEP': '09', 'OCTOBER': '10', 'OCT': '10', 'NOVEMBER': '11',
        'NOV': '11', 'DECEMBER': '12', 'DEC': '12', 'JANUARY': '01',
        'JAN': '01', 'FEBRUARY': '02', 'FEB': '02', 'MARCH': '03',
        'MAR': '03', 'APRIL': '04', 'APR': '04', 'MAY': '05'
    }
    return month_mapping.get(month_name.upper(), '00')

def get_month_term_index(month_name, term):
    """Get the month index within the term."""
    if term == '1':
        months = ['JUNE', 'JULY', 'AUG', 'SEP', 'OCT']
    else:  # term 2
        months = ['NOV', 'DEC', 'JAN', 'FEB']
    
    for i, m in enumerate(months, 1):
        if month_name.upper().startswith(m):
            return f"{i:02d}"
    return "01"  # Default

def replace_placeholder(text, replacements):
    """Replace placeholders in text."""
    for placeholder, value in replacements.items():
        text = text.replace(placeholder, value)
    return text

def create_month_word_file(month_name, data_rows, columns, file_info, temp_folder):
    """Create a Word document for a single month."""
    # Create a new document from template
    doc = Document('executive_summary_template.docx')
    
    # Replace placeholders in the document
    month_num = get_month_number(month_name)
    term_month_index = get_month_term_index(month_name, file_info['term'])
    
    # Create replacements dictionary
    replacements = {
        "MONTH-JUNE": f"MONTH-{month_name}",
        "2023-2024-06/JES/ECO/01": f"{file_info['year_range']}-{month_num}/JES/ECO/{term_month_index}",
        "DI/F-ES/00": f"DI/F-ES/{file_info['standard']}"
    }
    
    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)
    
    # Find the main table to populate
    table = None
    for potential_table in doc.tables:
        if len(potential_table.rows) > 0:
            header_row = potential_table.rows[0]
            for cell in header_row.cells:
                if "XI" in cell.text or "XII" in cell.text:
                    table = potential_table
                    break
            if table:
                break
    
    if not table:
        logger.error(f"No table found with XI or XII headers in template")
        return None
    
    # Determine column offset based on standard (FYJC/SYJC)
    col_offset = 0
    if file_info['original_std'] == 'SYJC':  # Use XII columns (columns 5-7)
        col_offset = 3
    
    # Fill table with data
    row_idx = 1  # Start after header row
    for data_row in data_rows:
        if row_idx >= len(table.rows):
            table.add_row()
        
        table_row = table.rows[row_idx]
        
        # Check if we have enough cells
        if len(table_row.cells) >= 8:
            # SR.NO.
            if len(data_row) > 0:
                table_row.cells[0].text = data_row[0]
            
            # INITIALS
            if len(data_row) > 1:
                table_row.cells[1].text = data_row[1]
            
            # Fill in ALOTTED, ENGAGED, GAP
            try:
                # ALLOTTED - Column index 2 for XI, 5 for XII
                if 'ALLOTTED' in columns and len(data_row) > columns['ALLOTTED']:
                    table_row.cells[2 + col_offset].text = data_row[columns['ALLOTTED']]
                
                # ENGAGED - Column index 3 for XI, 6 for XII
                if 'ENGAGED' in columns and len(data_row) > columns['ENGAGED']:
                    table_row.cells[3 + col_offset].text = data_row[columns['ENGAGED']]
                
                # GAP - Column index 4 for XI, 7 for XII
                if 'GAP' in columns and len(data_row) > columns['GAP']:
                    # Fix for gap value - E-Add is +value but we want just +value
                    gap_value = data_row[columns['GAP']]
                    if gap_value.startswith('+'):
                        table_row.cells[4 + col_offset].text = gap_value
                    else:
                        table_row.cells[4 + col_offset].text = f"+{gap_value}"
            except IndexError as e:
                logger.error(f"Index error when filling table: {e}")
        
        row_idx += 1
    
    # Save the document to temp folder
    output_filename = f"{file_info['original_std']}_{month_name}_{file_info['year_range']}.docx"
    output_path = os.path.join(temp_folder, output_filename)
    doc.save(output_path)
    logger.info(f"Created Word document for {month_name}: {output_path}")
    
    return output_path

def merge_word_documents(files_to_merge, output_file):
    """Merge multiple Word documents into one."""
    if not files_to_merge:
        logger.error("No files to merge")
        return False
    
    # Create a new document to hold the merged content
    merged_doc = Document()
    
    for i, file_path in enumerate(files_to_merge):
        logger.info(f"Merging file: {file_path}")
        
        # Add a page break before each file (except the first one)
        if i > 0:
            merged_doc.add_page_break()
        
        # Load the source document
        src_doc = Document(file_path)
        
        # Copy all paragraphs
        for paragraph in src_doc.paragraphs:
            # Skip empty paragraphs
            if not paragraph.text.strip():
                continue
            
            # Add paragraph to merged document
            p = merged_doc.add_paragraph()
            # Copy the runs with their formatting
            for run in paragraph.runs:
                merged_run = p.add_run(run.text)
                merged_run.bold = run.bold
                merged_run.italic = run.italic
                merged_run.underline = run.underline
                if run.font.color:
                    merged_run.font.color.rgb = run.font.color.rgb
                if run.font.size:
                    merged_run.font.size = run.font.size
        
        # Copy all tables
        for table in src_doc.tables:
            # Create new table with same dimensions
            new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            
            # Apply table style if available
            if table.style:
                new_table.style = table.style
            
            # Copy content and formatting cell by cell
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                        # Copy text and paragraphs with formatting
                        for paragraph in cell.paragraphs:
                            if not new_table.rows[i].cells[j].paragraphs:
                                cell_para = new_table.rows[i].cells[j].add_paragraph()
                            else:
                                cell_para = new_table.rows[i].cells[j].paragraphs[0]
                            
                            for run in paragraph.runs:
                                cell_run = cell_para.add_run(run.text)
                                cell_run.bold = run.bold
                                cell_run.italic = run.italic
                                cell_run.underline = run.underline
                                if run.font.color:
                                    cell_run.font.color.rgb = run.font.color.rgb
                                if run.font.size:
                                    cell_run.font.size = run.font.size
    
    # Save the merged document
    merged_doc.save(output_file)
    logger.info(f"Merged document saved to: {output_file}")
    return True

def create_word_documents(csv_path, output_folder="output_word_files", temp_folder="temp_word_files"):
    """Create Word documents from Excel data, one per month."""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    if not os.path.exists(temp_folder):
        os.makedirs(temp_folder)
    else:
        # Clean temp folder
        for file in os.listdir(temp_folder):
            file_path = os.path.join(temp_folder, file)
            try:
                if os.path.isfile(file_path):
                    os.unlink(file_path)
            except Exception as e:
                logger.error(f"Error deleting file {file_path}: {e}")
    
    file_info = parse_filename(csv_path)
    if not file_info:
        logger.error(f"Could not parse information from filename: {csv_path}")
        return False
    
    logger.info(f"Processing file: {csv_path}")
    logger.info(f"File info: {file_info}")
    
    month_files = []
    
    with open(csv_path, 'r', newline='') as csvfile:
        reader = csv.reader(csvfile)
        
        # First row contains months
        header_row = next(reader)
        
        # Second row contains column labels (ALOTTED, E-Act, etc.)
        column_labels = next(reader)
        
        # Map the column indices for each month
        month_columns = {}
        current_month = None
        
        for i, label in enumerate(header_row):
            # Check if this is a month column
            if label and label.upper() in ['JUNE', 'JULY', 'AUG', 'SEP', 'OCT', 'NOV', 'DEC', 'JAN', 'FEB', 'MAR', 'APR', 'MAY']:
                current_month = label.upper()
                month_columns[current_month] = {}
            
            if current_month and i < len(column_labels):
                col_label = column_labels[i].upper() if i < len(column_labels) else ""
                if col_label == 'ALOTTED':
                    month_columns[current_month]['ALLOTTED'] = i
                elif col_label == 'E-ACT':
                    month_columns[current_month]['ENGAGED'] = i
                elif col_label == 'E-ADD':
                    month_columns[current_month]['GAP'] = i
        
        logger.info(f"Month columns mapping: {month_columns}")
        
        # Store all data rows for processing
        all_data = []
        # Skip first 2 rows (already read)
        for row in reader:
            if row and len(row) > 0 and row[0].strip():
                if row[0].strip().upper() == 'TOTAL':
                    break
                all_data.append(row)
        
        # Process each month
        for month_name, columns in month_columns.items():
            if not all(key in columns for key in ['ALLOTTED', 'ENGAGED', 'GAP']):
                logger.warning(f"Skipping month {month_name} - missing required columns: {columns}")
                continue
            
            # Create document for this month
            month_file = create_month_word_file(month_name, all_data, columns, file_info, temp_folder)
            if month_file:
                month_files.append(month_file)
    
    # Merge all month files into one document
    if month_files:
        output_file = os.path.join(output_folder, f"{file_info['original_std']}_{file_info['year_range']}_term{file_info['term']}_all_months.docx")
        merge_word_documents(month_files, output_file)
        
        # Delete temp files if successful
        # for file in month_files:
        #     try:
        #         os.remove(file)
        #     except Exception as e:
        #         logger.error(f"Error deleting temp file {file}: {e}")
    else:
        logger.error("No month files were created to merge")
    
    return True

def process_excel_files(folder_path="excel_copies", output_folder="output_word_files", temp_folder="temp_word_files"):
    """Process all Excel files in the given folder."""
    if not os.path.exists(folder_path):
        logger.error(f"Folder not found: {folder_path}")
        return
    
    for filename in os.listdir(folder_path):
        if filename.endswith('.csv') and filename.startswith('iso_excel_'):
            file_path = os.path.join(folder_path, filename)
            logger.info(f"Processing: {filename}")
            create_word_documents(file_path, output_folder, temp_folder)

# Example usage
if __name__ == "__main__":
    logger.info("Starting Excel to Word conversion")
    process_excel_files()
    logger.info("Conversion completed")
