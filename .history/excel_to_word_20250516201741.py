import os
import csv
import re
import logging
from docx import Document
from docx.enum.section import WD_SECTION_START

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def parse_filename(filename):
    """Extract information from the filename."""
    # Expected format: iso_excel_YYYY-YYYY_termN_STD.csv
    pattern = r'iso_excel_(\d{4}-\d{4})_term(\d+)_(\w+)\.csv'
    match = re.match(pattern, os.path.basename(filename))
    
    if match:
        year_range, term, std = match.groups()
        # Convert STD to actual standard (FYJC=XI and SYJC=XII)
        std_mapping = {"FYJC": "XI", "SYJC": "XII"}
        standard = std_mapping.get(std, std)
        
        return {
            'year_range': year_range,
            'term': term,
            'standard': standard,
            'original_std': std
        }
    return None

def get_month_number(month_name):
    """Convert month name to number."""
    month_mapping = {
        'JUNE': '06', 'JULY': '07', 'AUG': '08', 'SEPTEMBER': '09', 
        'SEP': '09', 'OCTOBER': '10', 'OCT': '10', 'NOVEMBER': '11',
        'NOV': '11', 'DECEMBER': '12', 'DEC': '12', 'JANUARY': '01',
        'JAN': '01', 'FEBRUARY': '02', 'FEB': '02', 'MARCH': '03',
        'MAR': '03', 'APRIL': '04', 'APR': '04', 'MAY': '05'
    }
    return month_mapping.get(month_name.upper(), '00')

def get_month_term_index(month_name, term):
    """Get the month index within the term."""
    if term == '1':
        months = ['JUNE', 'JULY', 'AUG', 'SEP', 'OCT']
    else:  # term 2
        months = ['NOV', 'DEC', 'JAN', 'FEB']
    
    for i, m in enumerate(months, 1):
        if month_name.upper().startswith(m):
            return f"{i:02d}"
    return "01"  # Default

def update_placeholder_text(paragraph, replacements):
    """Update text in paragraph with placeholder replacements."""
    text = paragraph.text
    original_text = text
    
    for placeholder, value in replacements.items():
        if placeholder in text:
            text = text.replace(placeholder, value)
    
    # Only update if changes were made to avoid formatting loss
    if text != original_text:
        paragraph.text = text

def create_word_document(csv_path, output_folder="output_word_files"):
    """Create a single Word document with all months from Excel data."""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    file_info = parse_filename(csv_path)
    if not file_info:
        logger.error(f"Could not parse information from filename: {csv_path}")
        return False
    
    logger.info(f"Processing file: {csv_path}")
    logger.info(f"File info: {file_info}")
    
    # Create a single document for all months
    doc = Document('executive_summary_template.docx')
    
    with open(csv_path, 'r', newline='') as csvfile:
        reader = csv.reader(csvfile)
        
        # First row contains months
        header_row = next(reader)
        
        # Second row contains column labels
        column_labels = next(reader)
        
        # Map the column indices for each month
        month_columns = {}
        current_month = None
        
        for i, label in enumerate(header_row):
            if label and label.upper() in ['JUNE', 'JULY', 'AUG', 'SEP', 'OCT', 'NOV', 'DEC', 'JAN', 'FEB', 'MAR', 'APR', 'MAY']:
                current_month = label.upper()
                month_columns[current_month] = {}
            
            if current_month and i < len(column_labels):
                col_label = column_labels[i].upper() if i < len(column_labels) else ""
                if col_label == 'ALOTTED':
                    month_columns[current_month]['ALLOTTED'] = i
                elif col_label == 'E-ACT':
                    month_columns[current_month]['ENGAGED'] = i
                elif col_label == 'E-ADD':
                    month_columns[current_month]['GAP'] = i
        
        logger.info(f"Month columns mapping: {month_columns}")
        
        # Store data for each month
        month_data = {}
        for month_name in month_columns.keys():
            # Reset the CSV reader to read data rows
            csvfile.seek(0)
            next(reader)  # Skip header row
            next(reader)  # Skip column labels
            
            rows = []
            for data_row in reader:
                if not data_row or len(data_row) < 2 or not data_row[0].strip():
                    continue  # Skip empty rows
                
                if data_row[0].strip().upper() == 'TOTAL':
                    break  # Stop at total row
                
                rows.append(data_row)
            
            month_data[month_name] = rows
        
        # Process the first month first
        if month_columns:
            first_month = next(iter(month_columns))
            columns = month_columns[first_month]
            
            if 'ALLOTTED' in columns and 'ENGAGED' in columns and 'GAP' in columns:
                # Replace placeholders in the document
                month_num = get_month_number(first_month)
                term_month_index = get_month_term_index(first_month, file_info['term'])
                
                # Create placeholder replacements
                replacements = {
                    "{{year}}": file_info['year_range'],
                    "{{act_mon}}": month_num,
                    "{{term_mon}}": term_month_index,
                    "{{month}}": first_month
                }
                
                # Update paragraphs
                for paragraph in doc.paragraphs:
                    update_placeholder_text(paragraph, replacements)
                
                # Find the main table and fill it
                if doc.tables:
                    for table in doc.tables:
                        # Check if this is the target table by looking for column headers
                        is_target_table = False
                        has_standard_col = False
                        
                        if len(table.rows) > 0:
                            for cell in table.rows[0].cells:
                                if "XI" in cell.text or "XII" in cell.text:
                                    is_target_table = True
                                    break
                        
                        if not is_target_table:
                            continue
                        
                        # Fill the table with data
                        row_idx = 1  # Skip header row
                        
                        for data_row in month_data[first_month]:
                            # Add a new row if needed
                            if row_idx >= len(table.rows):
                                table.add_row()
                            
                            # Get the row in the table
                            table_row = table.rows[row_idx]
                            
                            # Fill in data
                            if len(table_row.cells) >= 8:  # Ensure the row has enough cells
                                # SR.NO.
                                if len(data_row) > 0 and data_row[0]:
                                    table_row.cells[0].text = data_row[0]
                                
                                # INITIALS
                                if len(data_row) > 1 and data_row[1]:
                                    table_row.cells[1].text = data_row[1]
                                
                                # Determine which column section to use (XI or XII)
                                # For SYJC use XII section
                                col_offset = 5 if file_info['original_std'] == 'SYJC' else 2
                                
                                # ALLOTTED
                                if len(data_row) > columns['ALLOTTED']:
                                    try:
                                        table_row.cells[col_offset].text = data_row[columns['ALLOTTED']]
                                    except IndexError:
                                        logger.error(f"Index error for ALLOTTED cell. Row: {row_idx}, Col: {col_offset}, Cells: {len(table_row.cells)}")
                                
                                # ENGAGED
                                if len(data_row) > columns['ENGAGED']:
                                    try:
                                        table_row.cells[col_offset + 1].text = data_row[columns['ENGAGED']]
                                    except IndexError:
                                        logger.error(f"Index error for ENGAGED cell. Row: {row_idx}, Col: {col_offset + 1}")
                                
                                # GAP
                                if len(data_row) > columns['GAP']:
                                    try:
                                        table_row.cells[col_offset + 2].text = data_row[columns['GAP']]
                                    except IndexError:
                                        logger.error(f"Index error for GAP cell. Row: {row_idx}, Col: {col_offset + 2}")
                            
                            row_idx += 1
        
        # Process the remaining months, each with a new page
        first_page_done = True
        for month_name, columns in list(month_columns.items())[1:]:
            if 'ALLOTTED' not in columns or 'ENGAGED' not in columns or 'GAP' not in columns:
                logger.warning(f"Skipping month {month_name} due to missing columns")
                continue
            
            # Add a page break
            doc.add_page_break()
            
            # Copy the original content from the template
            template = Document('executive_summary_template.docx')
            
            # Add the template content
            for i, paragraph in enumerate(template.paragraphs):
                if i >= len(doc.paragraphs):
                    doc.add_paragraph()
                
                p = doc.paragraphs[-1]
                p.text = paragraph.text
                
                # Apply replacements to the new paragraph
                month_num = get_month_number(month_name)
                term_month_index = get_month_term_index(month_name, file_info['term'])
                
                replacements = {
                    "{{year}}": file_info['year_range'],
                    "{{act_mon}}": month_num,
                    "{{term_mon}}": term_month_index,
                    "{{month}}": month_name
                }
                
                update_placeholder_text(p, replacements)
            
            # Add tables from template
            for template_table in template.tables:
                # Add a table to the document
                new_table = doc.add_table(rows=len(template_table.rows), cols=len(template_table.columns))
                
                # Copy table style
                if template_table.style:
                    new_table.style = template_table.style
                
                # Copy table cells
                for i, row in enumerate(template_table.rows):
                    for j, cell in enumerate(row.cells):
                        if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                            new_table.rows[i].cells[j].text = cell.text
                
                # Check if this is a data table
                is_target_table = False
                for cell in template_table.rows[0].cells:
                    if "XI" in cell.text or "XII" in cell.text:
                        is_target_table = True
                        break
                
                if is_target_table:
                    # Fill the new table with data
                    row_idx = 1  # Skip header row
                    
                    for data_row in month_data[month_name]:
                        # Add a new row if needed
                        if row_idx >= len(new_table.rows):
                            new_table.add_row()
                        
                        # Get the row in the table
                        table_row = new_table.rows[row_idx]
                        
                        # Fill in data
                        if len(table_row.cells) >= 8:  # Ensure the row has enough cells
                            # SR.NO.
                            if len(data_row) > 0 and data_row[0]:
                                table_row.cells[0].text = data_row[0]
                            
                            # INITIALS
                            if len(data_row) > 1 and data_row[1]:
                                table_row.cells[1].text = data_row[1]
                            
                            # Determine which column section to use (XI or XII)
                            col_offset = 5 if file_info['original_std'] == 'SYJC' else 2
                            
                            # ALLOTTED
                            if len(data_row) > columns['ALLOTTED']:
                                try:
                                    table_row.cells[col_offset].text = data_row[columns['ALLOTTED']]
                                except IndexError:
                                    logger.error(f"Index error for ALLOTTED cell in month {month_name}")
                            
                            # ENGAGED
                            if len(data_row) > columns['ENGAGED']:
                                try:
                                    table_row.cells[col_offset + 1].text = data_row[columns['ENGAGED']]
                                except IndexError:
                                    logger.error(f"Index error for ENGAGED cell in month {month_name}")
                            
                            # GAP
                            if len(data_row) > columns['GAP']:
                                try:
                                    table_row.cells[col_offset + 2].text = data_row[columns['GAP']]
                                except IndexError:
                                    logger.error(f"Index error for GAP cell in month {month_name}")
                        
                        row_idx += 1
                
                # Add some space after the table
                doc.add_paragraph()
    
    # Save the document
    output_filename = f"{file_info['original_std']}_{file_info['year_range']}_term{file_info['term']}_all_months.docx"
    output_path = os.path.join(output_folder, output_filename)
    doc.save(output_path)
    logger.info(f"Created Word document with all months: {output_path}")
    
    return True

def process_excel_files(folder_path="excel_copies", output_folder="output_word_files"):
    """Process all Excel files in the given folder."""
    if not os.path.exists(folder_path):
        logger.error(f"Folder not found: {folder_path}")
        return
    
    for filename in os.listdir(folder_path):
        if filename.endswith('.csv') and filename.startswith('iso_excel_'):
            file_path = os.path.join(folder_path, filename)
            logger.info(f"Processing: {filename}")
            create_word_document(file_path, output_folder)

# Example usage
if __name__ == "__main__":
    logger.info("Starting Excel to Word conversion")
    # Process all Excel files in the folder
    process_excel_files()
    logger.info("Conversion completed")
