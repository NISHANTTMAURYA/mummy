import os
import csv
import re
import logging
from docx import Document
from docx.enum.section import WD_SECTION_START

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def parse_filename(filename):
    """Extract information from the filename."""
    # Expected format: iso_excel_YYYY-YYYY_termN_STD.csv
    pattern = r'iso_excel_(\d{4}-\d{4})_term(\d+)_(\w+)\.csv'
    match = re.match(pattern, os.path.basename(filename))
    
    if match:
        year_range, term, std = match.groups()
        # Convert STD to actual standard (FYJC=XI and SYJC=XII)
        std_mapping = {"FYJC": "XI", "SYJC": "XII"}
        standard = std_mapping.get(std, std)
        
        return {
            'year_range': year_range,
            'term': term,
            'standard': standard,
            'original_std': std
        }
    return None

def get_month_number(month_name):
    """Convert month name to number."""
    month_mapping = {
        'JUNE': '06', 'JULY': '07', 'AUG': '08', 'SEPTEMBER': '09', 
        'SEP': '09', 'OCTOBER': '10', 'OCT': '10', 'NOVEMBER': '11',
        'NOV': '11', 'DECEMBER': '12', 'DEC': '12', 'JANUARY': '01',
        'JAN': '01', 'FEBRUARY': '02', 'FEB': '02', 'MARCH': '03',
        'MAR': '03', 'APRIL': '04', 'APR': '04', 'MAY': '05'
    }
    return month_mapping.get(month_name.upper(), '00')

def get_month_term_index(month_name, term):
    """Get the month index within the term."""
    if term == '1':
        months = ['JUNE', 'JULY', 'AUG', 'SEP', 'OCT']
    else:  # term 2
        months = ['NOV', 'DEC', 'JAN', 'FEB']
    
    for i, m in enumerate(months, 1):
        if month_name.upper().startswith(m):
            return f"{i:02d}"
    return "01"  # Default

def create_word_document(csv_path, output_folder="output_word_files"):
    """Create a single Word document with all months from Excel data."""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    file_info = parse_filename(csv_path)
    if not file_info:
        logger.error(f"Could not parse information from filename: {csv_path}")
        return False
    
    logger.info(f"Processing file: {csv_path}")
    logger.info(f"File info: {file_info}")
    
    # Create a new document from template
    template_doc = Document('executive_summary_template.docx')
    all_months_doc = Document()
    
    # Copy template styles to new document
    for style in template_doc.styles:
        if style.name not in all_months_doc.styles:
            try:
                all_months_doc.styles.add_style(style.name, style.type)
            except:
                pass
    
    with open(csv_path, 'r', newline='') as csvfile:
        reader = csv.reader(csvfile)
        
        # First row contains months
        header_row = next(reader)
        
        # Second row contains column labels
        column_labels = next(reader)
        
        # Map the column indices for each month
        month_columns = {}
        current_month = None
        
        for i, label in enumerate(header_row):
            if label and label.upper() in ['JUNE', 'JULY', 'AUG', 'SEP', 'OCT', 'NOV', 'DEC', 'JAN', 'FEB', 'MAR', 'APR', 'MAY']:
                current_month = label.upper()
                month_columns[current_month] = {}
            
            if current_month and i < len(column_labels):
                col_label = column_labels[i].upper() if i < len(column_labels) else ""
                if col_label == 'ALOTTED':
                    month_columns[current_month]['ALLOTTED'] = i
                elif col_label == 'E-ACT':
                    month_columns[current_month]['ENGAGED'] = i
                elif col_label == 'E-ADD':
                    month_columns[current_month]['GAP'] = i
        
        logger.info(f"Month columns mapping: {month_columns}")
        
        # Process each month
        first_page = True
        
        for month_name, columns in month_columns.items():
            logger.info(f"Processing month: {month_name} with columns: {columns}")
            if 'ALLOTTED' not in columns or 'ENGAGED' not in columns or 'GAP' not in columns:
                logger.warning(f"Missing required columns for month {month_name}, skipping. Found: {columns}")
                continue
            
            # Add a section break for each month except the first
            if not first_page:
                all_months_doc.add_section(WD_SECTION_START.NEW_PAGE)
            else:
                first_page = False
            
            # Copy template content for this month
            for element in template_doc.element.body:
                all_months_doc.element.body.append(element)
            
            # Replace placeholders in the document
            month_num = get_month_number(month_name)
            term_month_index = get_month_term_index(month_name, file_info['term'])
            
            # Get current section paragraphs
            section_paragraphs = all_months_doc.paragraphs[-len(template_doc.paragraphs):]
            
            for paragraph in section_paragraphs:
                # Replace month placeholder
                if "{{month}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{month}}", month_name.title())
                
                # Replace year and date code placeholders
                if "{{year}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{year}}", file_info['year_range'])
                if "{{act_mon}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{act_mon}}", month_num)
                if "{{term_mon}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{term_mon}}", term_month_index)
                if "JES/ECO/{{term_mon}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("JES/ECO/{{term_mon}}", f"JES/ECO/{term_month_index}")
                if "DI/F-ES/00" in paragraph.text:
                    paragraph.text = paragraph.text.replace("DI/F-ES/00", f"DI/F-ES/{file_info['standard']}")
            
            # Get tables for this section
            section_tables = all_months_doc.tables[-len(template_doc.tables):]
            
            # Process data for the tables
            for table in section_tables:
                # Check if this is our target table by looking for XI and XII columns
                is_target_table = False
                if len(table.rows) > 0:
                    for cell in table.rows[0].cells:
                        if "XI" in cell.text or "XII" in cell.text:
                            is_target_table = True
                            break
                
                if not is_target_table:
                    continue
                
                # Reset the CSV reader to read data rows
                csvfile.seek(0)
                next(reader)  # Skip header row
                next(reader)  # Skip column labels
                
                # Process each data row
                row_idx = 1  # Start after header row
                
                for data_row in reader:
                    if not data_row or len(data_row) < 2 or not data_row[0].strip():
                        continue  # Skip empty rows
                    
                    if data_row[0].strip().upper() == 'TOTAL':
                        break  # Stop at total row
                    
                    # Add a new row if needed
                    if row_idx >= len(table.rows):
                        table.add_row()
                    
                    # Get the row in the table
                    table_row = table.rows[row_idx]
                    
                    # Fill in data
                    if len(table_row.cells) >= 5:  # Ensure the row has enough cells
                        # SR.NO.
                        if len(data_row) > 0 and data_row[0]:
                            table_row.cells[0].text = data_row[0]
                        
                        # INITIALS
                        if len(data_row) > 1 and data_row[1]:
                            table_row.cells[1].text = data_row[1]
                        
                        # Determine which column section to use (XI or XII)
                        # For SYJC use XII section (columns offset by 3)
                        col_offset = 3 if file_info['original_std'] == 'SYJC' else 0
                        
                        try:
                            # ALLOTTED
                            if len(data_row) > columns['ALLOTTED']:
                                allotted_col = 2 + col_offset
                                if allotted_col < len(table_row.cells):
                                    table_row.cells[allotted_col].text = data_row[columns['ALLOTTED']]
                                else:
                                    logger.error(f"Cell index out of range: {allotted_col} >= {len(table_row.cells)}")
                            
                            # ENGAGED
                            if len(data_row) > columns['ENGAGED']:
                                engaged_col = 3 + col_offset
                                if engaged_col < len(table_row.cells):
                                    table_row.cells[engaged_col].text = data_row[columns['ENGAGED']]
                                else:
                                    logger.error(f"Cell index out of range: {engaged_col} >= {len(table_row.cells)}")
                            
                            # GAP
                            if len(data_row) > columns['GAP']:
                                gap_col = 4 + col_offset
                                if gap_col < len(table_row.cells):
                                    table_row.cells[gap_col].text = data_row[columns['GAP']]
                                else:
                                    logger.error(f"Cell index out of range: {gap_col} >= {len(table_row.cells)}")
                        except Exception as e:
                            logger.error(f"Error filling table: {e}")
                    
                    row_idx += 1
    
    # Save the document with all months
    output_filename = f"{file_info['original_std']}_{file_info['year_range']}_term{file_info['term']}_all_months.docx"
    output_path = os.path.join(output_folder, output_filename)
    all_months_doc.save(output_path)
    logger.info(f"Created Word document with all months: {output_path}")
    
    return True

def process_excel_files(folder_path="excel_copies", output_folder="output_word_files"):
    """Process all Excel files in the given folder."""
    if not os.path.exists(folder_path):
        logger.error(f"Folder not found: {folder_path}")
        return
    
    for filename in os.listdir(folder_path):
        if filename.endswith('.csv') and filename.startswith('iso_excel_'):
            file_path = os.path.join(folder_path, filename)
            logger.info(f"Processing: {filename}")
            create_word_document(file_path, output_folder)

# Example usage
if __name__ == "__main__":
    logger.info("Starting Excel to Word conversion")
    # Process a single file
    # create_word_document("excel_copies/iso_excel_2023-2024_term1_SYJC.csv")
    
    # Process all Excel files in the folder
    process_excel_files()
    logger.info("Conversion completed")
