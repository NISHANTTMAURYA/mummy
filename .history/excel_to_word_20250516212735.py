import os
import csv
import re
import logging
from docx import Document
import shutil
from copy import deepcopy

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def parse_filename(filename):
    """Extract information from the filename."""
    # Expected format: iso_excel_YYYY-YYYY_termN_STD.csv
    pattern = r'iso_excel_(\d{4}-\d{4})_term(\d+)_(\w+)\.csv'
    match = re.match(pattern, os.path.basename(filename))
    
    if match:
        year_range, term, std = match.groups()
        # Convert STD to actual standard (FYJC=XI and SYJC=XII)
        std_mapping = {"FYJC": "XI", "SYJC": "XII"}
        standard = std_mapping.get(std, std)
        
        return {
            'year_range': year_range,
            'term': term,
            'standard': standard,
            'original_std': std
        }
    return None

def get_month_number(month_name):
    """Convert month name to number."""
    month_mapping = {
        'JUNE': '06', 'JULY': '07', 'AUG': '08', 'SEPTEMBER': '09', 
        'SEP': '09', 'OCTOBER': '10', 'OCT': '10', 'NOVEMBER': '11',
        'NOV': '11', 'DECEMBER': '12', 'DEC': '12', 'JANUARY': '01',
        'JAN': '01', 'FEBRUARY': '02', 'FEB': '02', 'MARCH': '03',
        'MAR': '03', 'APRIL': '04', 'APR': '04', 'MAY': '05'
    }
    return month_mapping.get(month_name.upper(), '00')

def get_term_month_index(month_name, term):
    """Get the month index within the term."""
    if term == '1':
        months = ['JUNE', 'JULY', 'AUG', 'SEP', 'OCT']
    else:  # term 2
        months = ['NOV', 'DEC', 'JAN', 'FEB']
    
    for i, m in enumerate(months, 1):
        if month_name.upper().startswith(m):
            return f"{i:02d}"
    return "01"  # Default

def replace_placeholders_in_paragraph(paragraph, replacements):
    """Replace placeholders in paragraph text without affecting other formatting."""
    if not any(placeholder in paragraph.text for placeholder in replacements.keys()):
        return False
    
    # Create a new paragraph text with replacements
    text = paragraph.text
    for placeholder, value in replacements.items():
        if placeholder in text:
            text = text.replace(placeholder, value)
    
    # Replace text in the paragraph
    paragraph.text = text
    return True

def process_single_month(template_path, month_name, data_rows, columns, file_info, output_path):
    """Process a single month and save to output path."""
    try:
        # Create a new document from template
        doc = Document(template_path)
        
        # Replace placeholders in the document
        month_num = get_month_number(month_name)
        term_month_index = get_term_month_index(month_name, term=file_info['term'])
        
        # Create the replacements dictionary
        replacements = {
            "{{year}}": file_info['year_range'],
            "{{act_mon}}": month_num,
            "{{term_mon}}": term_month_index,
            "{{month}}": month_name,
            "ES/00": f"ES/{file_info['standard']}"
        }
        
        # Replace in paragraphs and table cells
        for paragraph in doc.paragraphs:
            replace_placeholders_in_paragraph(paragraph, replacements)
        
        # Check tables for placeholders and find data table
        target_table = None
        for table in doc.tables:
            if len(table.rows) > 0:
                # Look for the data table (one with XI or XII headers)
                has_std_header = False
                for cell in table.rows[0].cells:
                    if "XI" in cell.text or "XII" in cell.text:
                        has_std_header = True
                        target_table = table
                        break
                
                # Replace placeholders in all table cells
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_placeholders_in_paragraph(paragraph, replacements)
        
        if not target_table:
            logger.error("Could not find target table in template")
            return False
        
        # Determine offset based on whether it's FYJC (XI) or SYJC (XII)
        col_offset = 0
        if file_info['original_std'] == 'SYJC':  # Use XII columns
            col_offset = 3
        
        # Fill table data
        row_idx = 1  # Start from the first data row (after header)
        for data_row in data_rows:
            if not data_row or len(data_row) < 2:
                continue
            
            # Skip if we've hit the total row
            if data_row[0].strip().upper() == 'TOTAL':
                break
            
            if row_idx >= len(target_table.rows):
                if row_idx < 10:
                    target_table.add_row()
                else:
                    break
            
            try:
                row = target_table.rows[row_idx]
                
                # Fill SR.NO and INITIALS
                if len(data_row) > 0:
                    row.cells[0].text = data_row[0].strip()
                if len(data_row) > 1:
                    row.cells[1].text = data_row[1].strip()
                
                # Fill data in appropriate columns based on standard
                if 'ALLOTTED' in columns and len(data_row) > columns['ALLOTTED'] and 2 + col_offset < len(row.cells):
                    row.cells[2 + col_offset].text = data_row[columns['ALLOTTED']]
                
                if 'ENGAGED' in columns and len(data_row) > columns['ENGAGED'] and 3 + col_offset < len(row.cells):
                    row.cells[3 + col_offset].text = data_row[columns['ENGAGED']]
                
                if 'GAP' in columns and len(data_row) > columns['GAP'] and 4 + col_offset < len(row.cells):
                    gap_value = data_row[columns['GAP']]
                    if not gap_value.startswith('+') and gap_value.strip():
                        gap_value = f"+{gap_value}"
                    row.cells[4 + col_offset].text = gap_value
            except IndexError as e:
                logger.error(f"Index error when filling row {row_idx}: {e}")
            
            row_idx += 1
        
        # Save document
        doc.save(output_path)
        logger.info(f"Processed {month_name} and saved to {output_path}")
        return True
    except Exception as e:
        logger.error(f"Error processing {month_name}: {e}")
        return False

def combine_word_files(file_paths, output_path):
    """Combine Word files using direct section copy."""
    try:
        # Create a new document with all the content
        combined_doc = Document()
        
        for i, file_path in enumerate(file_paths):
            # Load source document
            src_doc = Document(file_path)
            
            # Copy each section (this preserves formatting and structure better)
            for j, section in enumerate(src_doc.sections):
                # Add a page break before each file except the first
                if i > 0 and j == 0:
                    combined_doc.add_page_break()
                
                # Create a new section in the target document
                if i > 0 or j > 0:
                    combined_doc.add_section()
                
                # Copy section properties
                target_section = combined_doc.sections[-1]
                target_section.start_type = section.start_type
                target_section.page_height = section.page_height
                target_section.page_width = section.page_width
                target_section.left_margin = section.left_margin
                target_section.right_margin = section.right_margin
                target_section.top_margin = section.top_margin
                target_section.bottom_margin = section.bottom_margin
                
                # Set up headers and footers (if needed)
                # Note: This is simplified and may need more detailed handling
                if section.header.is_linked_to_previous:
                    target_section.header.is_linked_to_previous = True
                
                if section.footer.is_linked_to_previous:
                    target_section.footer.is_linked_to_previous = True
            
            # Copy content (paragraphs and tables)
            for element in src_doc.element.body:
                combined_doc.element.body.append(element)
        
        # Save the combined document
        combined_doc.save(output_path)
        logger.info(f"Successfully combined {len(file_paths)} documents into {output_path}")
        return True
    except Exception as e:
        logger.error(f"Error combining documents: {e}")
        return False

def create_multi_month_document(csv_path, template_path, output_folder="output_word_files"):
    """Create a multi-month Word document from Excel data."""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    # Create temp folder for individual month files
    temp_folder = os.path.join(output_folder, "temp")
    if not os.path.exists(temp_folder):
        os.makedirs(temp_folder)
    
    file_info = parse_filename(csv_path)
    if not file_info:
        logger.error(f"Could not parse information from filename: {csv_path}")
        return False
    
    logger.info(f"Processing file: {csv_path}")
    logger.info(f"File info: {file_info}")
    
    month_files = []
    
    try:
        with open(csv_path, 'r', newline='') as csvfile:
            reader = csv.reader(csvfile)
            
            # First row contains months
            header_row = next(reader)
            
            # Second row contains column labels
            column_labels = next(reader)
            
            # Map column indices for each month
            month_columns = {}
            current_month = None
            
            for i, label in enumerate(header_row):
                if label and label.upper() in ['JUNE', 'JULY', 'AUG', 'SEP', 'OCT', 'NOV', 'DEC', 'JAN', 'FEB', 'MAR', 'APR', 'MAY']:
                    current_month = label.upper()
                    month_columns[current_month] = {}
                
                if current_month and i < len(column_labels):
                    col_label = column_labels[i].upper() if i < len(column_labels) else ""
                    if col_label == 'ALOTTED':
                        month_columns[current_month]['ALLOTTED'] = i
                    elif col_label == 'E-ACT':
                        month_columns[current_month]['ENGAGED'] = i
                    elif col_label == 'E-ADD':
                        month_columns[current_month]['GAP'] = i
            
            logger.info(f"Month columns mapping: {month_columns}")
            
            # Store all data rows
            all_data = []
            for row in reader:
                if row and len(row) > 0 and row[0].strip():
                    all_data.append(row)
                    if row[0].strip().upper() == 'TOTAL':
                        break
            
            # Process each month
            for month_name, columns in month_columns.items():
                if not all(key in columns for key in ['ALLOTTED', 'ENGAGED', 'GAP']):
                    logger.warning(f"Skipping month {month_name} - missing columns: {columns}")
                    continue
                
                # Create output file path for this month
                month_file = os.path.join(temp_folder, f"{file_info['original_std']}_{month_name}_{file_info['year_range']}.docx")
                
                # Process the month
                success = process_single_month(
                    template_path=template_path,
                    month_name=month_name,
                    data_rows=all_data,
                    columns=columns,
                    file_info=file_info,
                    output_path=month_file
                )
                
                if success:
                    month_files.append(month_file)
            
            # Combine all month files into one
            if month_files:
                output_file = os.path.join(output_folder, f"{file_info['original_std']}_{file_info['year_range']}_term{file_info['term']}_all_months.docx")
                combine_word_files(month_files, output_file)
            else:
                logger.error("No month files were created to combine")
    
    except Exception as e:
        logger.error(f"Error processing file {csv_path}: {e}")
        return False
    
    finally:
        # Clean up temp files
        for file_path in month_files:
            try:
                os.remove(file_path)
            except:
                pass
        try:
            os.rmdir(temp_folder)
        except:
            pass
    
    return True

def process_excel_files(excel_folder="excel_copies", template_path="executive_summary_template.docx", output_folder="output_word_files"):
    """Process all Excel files in the given folder."""
    if not os.path.exists(excel_folder):
        logger.error(f"Excel folder not found: {excel_folder}")
        return
    
    if not os.path.exists(template_path):
        logger.error(f"Template file not found: {template_path}")
        return
    
    for filename in os.listdir(excel_folder):
        if filename.endswith('.csv') and filename.startswith('iso_excel_'):
            file_path = os.path.join(excel_folder, filename)
            logger.info(f"Processing: {filename}")
            create_multi_month_document(file_path, template_path, output_folder)

# Example usage
if __name__ == "__main__":
    logger.info("Starting Excel to Word conversion")
    process_excel_files()
    logger.info("Conversion completed")
