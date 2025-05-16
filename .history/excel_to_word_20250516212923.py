import os
import csv
import re
import logging
from docx import Document
import shutil
from copy import deepcopy

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def parse_filename(filename):
    """Extract information from the filename."""
    # Expected format: iso_excel_YYYY-YYYY_termN_STD.csv
    pattern = r'iso_excel_(\d{4}-\d{4})_term(\d+)_(\w+)\.csv'
    match = re.match(pattern, os.path.basename(filename))
    
    if match:
        year_range, term, std = match.groups()
        # Convert STD to actual standard (FYJC=XI and SYJC=XII)
        std_mapping = {"FYJC": "XI", "SYJC": "XII"}
        standard = std_mapping.get(std, std)
        
        return {
            'year_range': year_range,
            'term': term,
            'standard': standard,
            'original_std': std
        }
    return None

def get_month_number(month_name):
    """Convert month name to number."""
    month_mapping = {
        'JUNE': '06', 'JULY': '07', 'AUG': '08', 'SEPTEMBER': '09', 
        'SEP': '09', 'OCTOBER': '10', 'OCT': '10', 'NOVEMBER': '11',
        'NOV': '11', 'DECEMBER': '12', 'DEC': '12', 'JANUARY': '01',
        'JAN': '01', 'FEBRUARY': '02', 'FEB': '02', 'MARCH': '03',
        'MAR': '03', 'APRIL': '04', 'APR': '04', 'MAY': '05'
    }
    return month_mapping.get(month_name.upper(), '00')

def get_term_month_index(month_name, term):
    """Get the month index within the term."""
    if term == '1':
        months = ['JUNE', 'JULY', 'AUG', 'SEP', 'OCT']
    else:  # term 2
        months = ['NOV', 'DEC', 'JAN', 'FEB']
    
    for i, m in enumerate(months, 1):
        if month_name.upper().startswith(m):
            return f"{i:02d}"
    return "01"  # Default

def replace_placeholders_in_paragraph(paragraph, replacements):
    """Replace placeholders in paragraph text without affecting other formatting."""
    if not any(placeholder in paragraph.text for placeholder in replacements.keys()):
        return False
    
    # Create a new paragraph text with replacements
    text = paragraph.text
    for placeholder, value in replacements.items():
        if placeholder in text:
            text = text.replace(placeholder, value)
    
    # Replace text in the paragraph
    paragraph.text = text
    return True

def process_single_month(template_path, month_name, data_rows, columns, file_info, output_path):
    """Process a single month and save to output path."""
    try:
        # Create a new document from template
        doc = Document(template_path)
        
        # Replace placeholders in the document
        month_num = get_month_number(month_name)
        term_month_index = get_term_month_index(month_name, term=file_info['term'])
        
        # Create the replacements dictionary
        replacements = {
            "{{year}}": file_info['year_range'],
            "{{act_mon}}": month_num,
            "{{term_mon}}": term_month_index,
            "{{month}}": month_name,
            "ES/00": f"ES/{file_info['standard']}"
        }
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            replace_placeholders_in_paragraph(paragraph, replacements)
        
        # Check tables for placeholders and find data table
        target_table = None
        for table in doc.tables:
            if len(table.rows) > 0:
                # Look for the data table (one with XI or XII headers)
                has_std_header = False
                for cell in table.rows[0].cells:
                    if "XI" in cell.text or "XII" in cell.text:
                        has_std_header = True
                        target_table = table
                        break
                
                # Replace placeholders in all table cells
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_placeholders_in_paragraph(paragraph, replacements)
        
        if not target_table:
            logger.error("Could not find target table in template")
            return False
        
        # Determine offset based on whether it's FYJC (XI) or SYJC (XII)
        col_offset = 0
        if file_info['original_std'] == 'SYJC':  # Use XII columns
            col_offset = 3
        
        # Fill table data
        row_idx = 1  # Start from the first data row (after header)
        for data_row in data_rows:
            if not data_row or len(data_row) < 2:
                continue
            
            # Skip if we've hit the total row
            if data_row[0].strip().upper() == 'TOTAL':
                break
            
            if row_idx >= len(target_table.rows):
                if row_idx < 10:
                    target_table.add_row()
                else:
                    break
            
            try:
                row = target_table.rows[row_idx]
                
                # Fill SR.NO and INITIALS
                if len(data_row) > 0:
                    row.cells[0].text = data_row[0].strip()
                if len(data_row) > 1:
                    row.cells[1].text = data_row[1].strip()
                
                # Fill data in appropriate columns based on standard
                if 'ALLOTTED' in columns and len(data_row) > columns['ALLOTTED'] and 2 + col_offset < len(row.cells):
                    row.cells[2 + col_offset].text = data_row[columns['ALLOTTED']]
                
                if 'ENGAGED' in columns and len(data_row) > columns['ENGAGED'] and 3 + col_offset < len(row.cells):
                    row.cells[3 + col_offset].text = data_row[columns['ENGAGED']]
                
                if 'GAP' in columns and len(data_row) > columns['GAP'] and 4 + col_offset < len(row.cells):
                    gap_value = data_row[columns['GAP']]
                    if not gap_value.startswith('+') and gap_value.strip():
                        gap_value = f"+{gap_value}"
                    row.cells[4 + col_offset].text = gap_value
            except IndexError as e:
                logger.error(f"Index error when filling row {row_idx}: {e}")
            
            row_idx += 1
        
        # Save document
        doc.save(output_path)
        logger.info(f"Processed {month_name} and saved to {output_path}")
        return True
    except Exception as e:
        logger.error(f"Error processing {month_name}: {e}")
        return False

def create_multi_month_document(csv_path, template_path, output_folder="output_word_files"):
    """Create individual month files and then provide instructions to manually combine them."""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    # Create month-specific folder for individual month files
    month_folder = os.path.join(output_folder, "month_files")
    if not os.path.exists(month_folder):
        os.makedirs(month_folder)
    
    file_info = parse_filename(csv_path)
    if not file_info:
        logger.error(f"Could not parse information from filename: {csv_path}")
        return False
    
    logger.info(f"Processing file: {csv_path}")
    logger.info(f"File info: {file_info}")
    
    month_files = []
    
    try:
        with open(csv_path, 'r', newline='') as csvfile:
            reader = csv.reader(csvfile)
            
            # First row contains months
            header_row = next(reader)
            
            # Second row contains column labels
            column_labels = next(reader)
            
            # Map column indices for each month
            month_columns = {}
            current_month = None
            
            for i, label in enumerate(header_row):
                if label and label.upper() in ['JUNE', 'JULY', 'AUG', 'SEP', 'OCT', 'NOV', 'DEC', 'JAN', 'FEB', 'MAR', 'APR', 'MAY']:
                    current_month = label.upper()
                    month_columns[current_month] = {}
                
                if current_month and i < len(column_labels):
                    col_label = column_labels[i].upper() if i < len(column_labels) else ""
                    if col_label == 'ALOTTED':
                        month_columns[current_month]['ALLOTTED'] = i
                    elif col_label == 'E-ACT':
                        month_columns[current_month]['ENGAGED'] = i
                    elif col_label == 'E-ADD':
                        month_columns[current_month]['GAP'] = i
            
            logger.info(f"Month columns mapping: {month_columns}")
            
            # Store all data rows
            all_data = []
            for row in reader:
                if row and len(row) > 0 and row[0].strip():
                    all_data.append(row)
                    if row[0].strip().upper() == 'TOTAL':
                        break
            
            # Process each month
            for month_name, columns in month_columns.items():
                if not all(key in columns for key in ['ALLOTTED', 'ENGAGED', 'GAP']):
                    logger.warning(f"Skipping month {month_name} - missing columns: {columns}")
                    continue
                
                # Create output file path for this month
                month_file = os.path.join(month_folder, f"{file_info['original_std']}_{month_name}_{file_info['year_range']}.docx")
                
                # Process the month
                success = process_single_month(
                    template_path=template_path,
                    month_name=month_name,
                    data_rows=all_data,
                    columns=columns,
                    file_info=file_info,
                    output_path=month_file
                )
                
                if success:
                    month_files.append(month_file)
            
            # Create a text file with instructions on how to manually combine the files
            if month_files:
                combined_filename = f"{file_info['original_std']}_{file_info['year_range']}_term{file_info['term']}_all_months"
                instruction_file = os.path.join(output_folder, f"{combined_filename}_INSTRUCTIONS.txt")
                
                with open(instruction_file, 'w') as f:
                    f.write("INSTRUCTIONS FOR COMBINING THE MONTH FILES\n")
                    f.write("===========================================\n\n")
                    f.write("Due to formatting issues when programmatically combining Word documents,\n")
                    f.write("please follow these steps to combine the month files manually:\n\n")
                    f.write("1. Open the first month file\n")
                    f.write("2. For each additional month file:\n")
                    f.write("   a. Place cursor at the end of the document\n")
                    f.write("   b. Insert -> Page Break\n")
                    f.write("   c. Insert -> Object -> Text from file\n")
                    f.write("   d. Select the next month file\n")
                    f.write("3. Save the combined document\n\n")
                    f.write("Individual month files are located in the 'month_files' folder:\n\n")
                    
                    for i, file_path in enumerate(month_files, 1):
                        f.write(f"{i}. {os.path.basename(file_path)}\n")
                
                logger.info(f"Created {len(month_files)} individual month files in the '{month_folder}' folder")
                logger.info(f"Instructions for manually combining files written to {instruction_file}")
            else:
                logger.error("No month files were created")
    
    except Exception as e:
        logger.error(f"Error processing file {csv_path}: {e}")
        return False
    
    return True

def process_excel_files(excel_folder="excel_copies", template_path="executive_summary_template.docx", output_folder="output_word_files"):
    """Process all Excel files in the given folder."""
    if not os.path.exists(excel_folder):
        logger.error(f"Excel folder not found: {excel_folder}")
        return
    
    if not os.path.exists(template_path):
        logger.error(f"Template file not found: {template_path}")
        return
    
    for filename in os.listdir(excel_folder):
        if filename.endswith('.csv') and filename.startswith('iso_excel_'):
            file_path = os.path.join(excel_folder, filename)
            logger.info(f"Processing: {filename}")
            create_multi_month_document(file_path, template_path, output_folder)

# Example usage
if __name__ == "__main__":
    logger.info("Starting Excel to Word conversion")
    process_excel_files()
    logger.info("Conversion completed")
