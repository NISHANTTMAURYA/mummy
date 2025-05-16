import os
import csv
import re
import logging
from docx import Document
import shutil
from copy import deepcopy

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def parse_filename(filename):
    """Extract information from the filename."""
    # Expected format: iso_excel_YYYY-YYYY_termN_STD.csv
    pattern = r'iso_excel_(\d{4}-\d{4})_term(\d+)_(\w+)\.csv'
    match = re.match(pattern, os.path.basename(filename))
    
    if match:
        year_range, term, std = match.groups()
        # Convert STD to actual standard (FYJC=XI and SYJC=XII)
        std_mapping = {"FYJC": "XI", "SYJC": "XII"}
        standard = std_mapping.get(std, std)
        
        return {
            'year_range': year_range,
            'term': term,
            'standard': standard,
            'original_std': std
        }
    return None

def get_month_number(month_name):
    """Convert month name to number."""
    month_mapping = {
        'JUNE': '06', 'JULY': '07', 'AUG': '08', 'SEPTEMBER': '09', 
        'SEP': '09', 'OCTOBER': '10', 'OCT': '10', 'NOVEMBER': '11',
        'NOV': '11', 'DECEMBER': '12', 'DEC': '12', 'JANUARY': '01',
        'JAN': '01', 'FEBRUARY': '02', 'FEB': '02', 'MARCH': '03',
        'MAR': '03', 'APRIL': '04', 'APR': '04', 'MAY': '05'
    }
    return month_mapping.get(month_name.upper(), '00')

def get_month_term_index(month_name, term):
    """Get the month index within the term."""
    if term == '1':
        months = ['JUNE', 'JULY', 'AUG', 'SEP', 'OCT']
    else:  # term 2
        months = ['NOV', 'DEC', 'JAN', 'FEB']
    
    for i, m in enumerate(months, 1):
        if month_name.upper().startswith(m):
            return f"{i:02d}"
    return "01"  # Default

def replace_placeholder(text, replacements):
    """Replace placeholders in text."""
    for placeholder, value in replacements.items():
        text = text.replace(placeholder, value)
    return text

def create_month_word_file(month_name, data_rows, columns, file_info, temp_folder):
    """Create a Word document for a single month."""
    try:
        # Create a new document from template - DO NOT MODIFY THE TEMPLATE, just copy it
        doc = Document('executive_summary_template.docx')
        
        # Locate the data table
        target_table = None
        for table in doc.tables:
            if len(table.rows) > 0:
                for cell in table.rows[0].cells:
                    if "XI" in cell.text or "XII" in cell.text:
                        target_table = table
                        break
                if target_table:
                    break
        
        if not target_table:
            logger.error("Could not find target table in template")
            return None
        
        # Determine offset based on whether it's FYJC (XI) or SYJC (XII)
        col_offset = 0
        if file_info['original_std'] == 'SYJC':  # Use XII columns
            col_offset = 3
        
        # Fill the table data only - DO NOT MODIFY ANY OTHER CONTENT
        row_idx = 1  # Start from the first data row (after header)
        for data_row in data_rows:
            if not data_row or len(data_row) < 2:
                continue
            
            # Skip if row_idx is beyond the table or if we've hit the total row
            if data_row[0].strip().upper() == 'TOTAL':
                break
            
            if row_idx >= len(target_table.rows):
                if row_idx < 10:  # Only add rows up to a reasonable limit
                    target_table.add_row()
                else:
                    break
            
            try:
                row = target_table.rows[row_idx]
                
                # Only fill SR.NO and INITIALS if they're empty
                if not row.cells[0].text.strip():
                    row.cells[0].text = data_row[0].strip()
                if not row.cells[1].text.strip() and len(data_row) > 1:
                    row.cells[1].text = data_row[1].strip()
                
                # Fill data in the appropriate columns (XI or XII) based on the file standard
                # Add data for allotted/engaged/gap columns if the column exists
                if 'ALLOTTED' in columns and len(data_row) > columns['ALLOTTED'] and 2 + col_offset < len(row.cells):
                    row.cells[2 + col_offset].text = data_row[columns['ALLOTTED']]
                
                if 'ENGAGED' in columns and len(data_row) > columns['ENGAGED'] and 3 + col_offset < len(row.cells):
                    row.cells[3 + col_offset].text = data_row[columns['ENGAGED']]
                
                if 'GAP' in columns and len(data_row) > columns['GAP'] and 4 + col_offset < len(row.cells):
                    gap_value = data_row[columns['GAP']]
                    # Ensure gap value has + prefix
                    if not gap_value.startswith('+') and gap_value.strip():
                        gap_value = f"+{gap_value}"
                    row.cells[4 + col_offset].text = gap_value
            except IndexError as e:
                logger.error(f"Index error when filling row {row_idx}: {e}")
            
            row_idx += 1
        
        # Save document to temp folder
        output_filename = f"{file_info['original_std']}_{month_name}_{file_info['year_range']}.docx"
        output_path = os.path.join(temp_folder, output_filename)
        doc.save(output_path)
        logger.info(f"Created Word document for {month_name}: {output_path}")
        
        return output_path
    except Exception as e:
        logger.error(f"Error creating Word file for {month_name}: {e}")
        return None

def append_documents(files_to_merge, output_file):
    """Append documents while preserving the exact template format."""
    if not files_to_merge:
        logger.error("No files to merge")
        return False
    
    try:
        # Use the first document as the base
        main_doc = Document(files_to_merge[0])
        
        # For each subsequent document, add a page break and copy the entire document content
        for i, file_path in enumerate(files_to_merge[1:], 1):
            logger.info(f"Appending document {i+1}/{len(files_to_merge)}: {file_path}")
            
            # Add a page break
            main_doc.add_page_break()
            
            # Open the document to append
            src_doc = Document(file_path)
            
            # Copy all paragraphs
            for paragraph in src_doc.paragraphs:
                main_doc.add_paragraph(paragraph.text)
            
            # Copy all tables - tables are more complex
            for table in src_doc.tables:
                # Create a new table with the same dimensions
                if len(table.rows) > 0:
                    num_cols = len(table.rows[0].cells)
                    new_table = main_doc.add_table(rows=len(table.rows), cols=num_cols)
                    
                    # Copy table style
                    if hasattr(table, 'style') and table.style:
                        new_table.style = table.style
                    
                    # Copy table content cell by cell
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                                new_table.rows[i].cells[j].text = cell.text
        
        # Save the merged document
        main_doc.save(output_file)
        logger.info(f"Successfully merged {len(files_to_merge)} documents to: {output_file}")
        return True
    except Exception as e:
        logger.error(f"Error appending documents: {e}")
        return False

def create_word_documents(csv_path, output_folder="output_word_files", temp_folder="temp_word_files"):
    """Create Word documents from Excel data, one per month."""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    if not os.path.exists(temp_folder):
        os.makedirs(temp_folder)
    else:
        # Clean temp folder
        for file in os.listdir(temp_folder):
            file_path = os.path.join(temp_folder, file)
            try:
                if os.path.isfile(file_path):
                    os.unlink(file_path)
            except Exception as e:
                logger.error(f"Error deleting file {file_path}: {e}")
    
    file_info = parse_filename(csv_path)
    if not file_info:
        logger.error(f"Could not parse information from filename: {csv_path}")
        return False
    
    logger.info(f"Processing file: {csv_path}")
    logger.info(f"File info: {file_info}")
    
    month_files = []
    
    with open(csv_path, 'r', newline='') as csvfile:
        reader = csv.reader(csvfile)
        
        # First row contains months
        header_row = next(reader)
        
        # Second row contains column labels (ALOTTED, E-Act, etc.)
        column_labels = next(reader)
        
        # Map the column indices for each month
        month_columns = {}
        current_month = None
        
        for i, label in enumerate(header_row):
            # Check if this is a month column
            if label and label.upper() in ['JUNE', 'JULY', 'AUG', 'SEP', 'OCT', 'NOV', 'DEC', 'JAN', 'FEB', 'MAR', 'APR', 'MAY']:
                current_month = label.upper()
                month_columns[current_month] = {}
            
            if current_month and i < len(column_labels):
                col_label = column_labels[i].upper() if i < len(column_labels) else ""
                if col_label == 'ALOTTED':
                    month_columns[current_month]['ALLOTTED'] = i
                elif col_label == 'E-ACT':
                    month_columns[current_month]['ENGAGED'] = i
                elif col_label == 'E-ADD':
                    month_columns[current_month]['GAP'] = i
        
        logger.info(f"Month columns mapping: {month_columns}")
        
        # Store all data rows for processing
        all_data = []
        # Skip first 2 rows (already read)
        for row in reader:
            if row and len(row) > 0 and row[0].strip():
                all_data.append(row)
                if row[0].strip().upper() == 'TOTAL':
                    break
        
        # Process each month
        for month_name, columns in month_columns.items():
            if not all(key in columns for key in ['ALLOTTED', 'ENGAGED', 'GAP']):
                logger.warning(f"Skipping month {month_name} - missing required columns: {columns}")
                continue
            
            # Create document for this month
            month_file = create_month_word_file(month_name, all_data, columns, file_info, temp_folder)
            if month_file:
                month_files.append(month_file)
    
    # Merge all month files into one document
    if month_files:
        output_file = os.path.join(output_folder, f"{file_info['original_std']}_{file_info['year_range']}_term{file_info['term']}_all_months.docx")
        append_documents(month_files, output_file)
    else:
        logger.error("No month files were created to merge")
    
    return True

def process_excel_files(folder_path="excel_copies", output_folder="output_word_files", temp_folder="temp_word_files"):
    """Process all Excel files in the given folder."""
    if not os.path.exists(folder_path):
        logger.error(f"Folder not found: {folder_path}")
        return
    
    for filename in os.listdir(folder_path):
        if filename.endswith('.csv') and filename.startswith('iso_excel_'):
            file_path = os.path.join(folder_path, filename)
            logger.info(f"Processing: {filename}")
            create_word_documents(file_path, output_folder, temp_folder)

# Example usage
if __name__ == "__main__":
    logger.info("Starting Excel to Word conversion")
    process_excel_files()
    logger.info("Conversion completed")
